# -*- coding: utf-8 -*-
"""生成云端模型选型文档（当前口径：Qwen2.5 系列）。

当前方案（2026-08 执行版任务书）：
  边缘端：Qwen2.5-1.5B Q4_K_M（质量优先）+ Qwen2.5-0.5B Q4_K_M（低内存兜底）
  云端：  Qwen2.5-14B（vLLM，待接入；mock 先行）
早期 Qwen3-30B-A3B/Qwen3-4B 方案已被实测数据取代，本脚本生成当前口径文档。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 默认样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(3)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cjk(run, name='微软雅黑'):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_heading(text, level=1, color='1A56C4'):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_cjk(r)
        r.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0,2,4)])
    return h

def add_para(text, bold=False, color=None, indent=False, size=10):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    set_cjk(r)
    return p

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def make_table(headers, widths, rows_data, highlight_keywords=('✅', '❌')):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text, w in zip(table.rows[0].cells, headers, widths):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cjk(r)
        set_cell_shading(cell, '1A56C4')
    for row_data in rows_data:
        row = table.add_row()
        for cell, text, w in zip(row.cells, row_data, widths):
            cell.width = w
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(9)
            set_cjk(r)
            if '✅' in text:
                r.font.color.rgb = RGBColor(0x00, 0x7a, 0x33)
                r.font.bold = True
            elif '❌' in text:
                r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
            elif '⚠️' in text:
                r.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
    return table

# ===== 封面 =====
for _ in range(5):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('云边协同 LLM 模型选型方案\n（Qwen2.5 双路径 + 云端 14B）')
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
set_cjk(r)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(20)
r = sub.add_run('智慧病房云边端协同系统 · 竞赛模型选型（2026-08 执行版）')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
set_cjk(r)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.paragraph_format.space_before = Pt(30)
r = info.add_run('边缘端：Qwen2.5-1.5B / 0.5B Q4（llama-cpp 本地推理）\n云端：Qwen2.5-14B（vLLM，独立 GPU 服务器）')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
set_cjk(r)

doc.add_page_break()

# ===== 一、选型概述 =====
add_heading('一、选型概述', 1)
add_para('本方案为智慧病房云边协同系统选择边缘与云端 LLM 模型，构成"边缘轻量模型 + 云端全量大模型"互补的云边协同推理架构。边缘端负责毫秒级语义增强、护理建议与离线自治决策；云端负责低置信度/高复杂度事件的二次研判。', indent=True)
add_para('边缘端采用 Qwen2.5 双路径策略：1.5B Q4_K_M 为质量优先路径，保留较强的事件语义理解与护理建议能力；0.5B Q4_K_M 为低内存路径，保证 8GB Jetson Orin Nano 上可运行。两条路径均已实测（Windows/x86 基线），最终主路径由 Jetson 实机数据裁决。', indent=True)
add_para('云端选用 Qwen2.5-14B（vLLM 部署），作为边缘端拿不准时的高置信度复核；云端服务已按契约实现 mock 消费者（cloud-llm-service），真实 14B 接入为上线前置项。', indent=True)

# ===== 二、边缘端模型对比 =====
add_heading('二、边缘端模型对比（Qwen2.5 双路径）', 1)

make_table(
    ['模型', '量化', '文件大小', '热身后 TTFT*', '峰值 RSS*', '定位', '选型'],
    [Cm(2.8), Cm(1.6), Cm(2.0), Cm(2.4), Cm(2.2), Cm(3.2), Cm(2.2)],
    [
        ['Qwen2.5-1.5B', 'Q4_K_M', '~1.04GB', '约 31.9ms', '约 1658MB', '质量优先：语义增强/护理建议更强', '✅ 主路径候选'],
        ['Qwen2.5-0.5B', 'Q4_K_M', '~469MB', '约 19.8ms', '约 516MB', '低内存：保证 8GB Jetson 可运行', '✅ 兜底路径'],
        ['Qwen3-4B', 'Q4_K_M', '~2.4GB', '—', '—', '早期候选：内存超标且 Jetson 跑不动', '❌ 已弃用'],
    ],
)

doc.add_paragraph()
add_para('* 数据来自 Windows/x86 实测（上下文 512、batch 128、8 线程、热身后）；'
         'Jetson Orin Nano 实机数据待 8/12 验收后替换本表。', size=9, color=RGBColor(0x88, 0x88, 0x88))

add_heading('1. 为什么保留两条路径', 2)
bullets = [
    '赛题硬指标"单次推理内存 ≤1.5GB"：1.5B 实测 1658MB 超标，0.5B 实测 516MB 达标',
    'Jetson 8GB 总内存需同时运行 YOLO-pose 视觉模型与容器/系统，0.5B 留足余量',
    '两条路径配置切换无需改代码：docker-compose.compact.yml 换模型文件 + 环境变量',
    '若 1.5B 在 Jetson 上超内存线，启用 compact 路径并保留质量/资源权衡的量化数据',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(b)
    r.font.size = Pt(10)
    set_cjk(r)

# ===== 三、云端模型选型 =====
add_heading('三、云端模型选型（Qwen2.5-14B）', 1)

make_table(
    ['模型', '参数量', '推理显存', '特点', '选型建议'],
    [Cm(2.5), Cm(2.0), Cm(2.2), Cm(4.5), Cm(4.6)],
    [
        ['Qwen2.5-14B', '14B', '~8GB (FP16)\n~5GB (INT4/AWQ)', '能力显著强于 1.5B，满足低置信度事件复核', '云端 ✅ 首选'],
        ['Qwen2.5-7B', '7B', '~4GB (FP16)\n~3GB (INT4)', '显存不足时的降级选择', '云端 备选'],
        ['Qwen3-30B-A3B', '30B\n(激活3B)', '~18GB (FP16)', 'MoE 高效，但需 ≥24GB 显存与授权', '云端 备选（资源足时）'],
    ],
)

doc.add_paragraph()
add_heading('1. 为什么选 14B 而非更大模型', 2)
bullets2 = [
    '与边缘 Qwen2.5-1.5B/0.5B 同属 Qwen2.5 系列，共享 tokenizer，蒸馏对齐成本低',
    '14B INT4/AWQ 约 5GB 显存，单张消费级/云端 GPU 即可部署，资源门槛可控',
    '30B MoE 需 ≥24GB 显存且需授权，超出当前团队可用资源，列为资源充足时的备选',
    '云边能力差距（14B vs 1.5B ≈ 9 倍）足以展示"边缘轻量 + 云端全量"互补效果',
]
for b in bullets2:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(b)
    r.font.size = Pt(10)
    set_cjk(r)

# ===== 四、部署方案 =====
add_heading('四、部署方案', 1)

add_heading('边缘端（Jetson Orin Nano 8GB / x86）', 2)
steps_edge = [
    '推理框架：llama-cpp-python（mmap 加载，上下文 512，batch 128，线程 8）',
    '模型文件：qwen2.5-1.5b-instruct-q4_k_m.gguf（质量路径）或 qwen2.5-0.5b-instruct-q4_k_m.gguf（compact 路径）',
    '主 Compose：docker compose up --build；compact 路径：docker compose -f docker-compose.yml -f docker-compose.compact.yml up --build',
    '环境变量：LLM_MODE=real、LLM_MODEL_PATH、LLM_N_CTX/LLM_N_BATCH/LLM_N_THREADS 等（见 .env.example）',
    '性能基准：python edge-agent/scripts/bench_jetson.py --model <gguf路径> --rounds 30',
]
for i, step in enumerate(steps_edge, 1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(step)
    r.font.size = Pt(10)
    set_cjk(r)

add_heading('云端（独立 GPU 服务器）', 2)
steps_cloud = [
    '推理框架：vLLM（推荐）',
    '模型下载：Qwen/Qwen2.5-14B-Instruct（AWQ/INT4 量化以降低显存）',
    '启动命令（vLLM）：python -m vllm.entrypoints.openai.api_server --model Qwen/Qwen2.5-14B-Instruct-AWQ --max-model-len 8192',
    '接入路径：cloud-llm-service 已实现 mock 消费者与响应契约（event_id/trace_id/judgment/confidence/advice/latency_ms），真实模型接入替换 LLMClient 后端即可',
    '未就绪前：以接口一致的 mock consumer 打通云边闭环，真实模型接入列为上线前置项（P7/P6）',
]
for i, step in enumerate(steps_cloud, 1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(step)
    r.font.size = Pt(10)
    set_cjk(r)

add_heading('云边协同推理流程', 2)
add_para('1. 边缘端 YOLO/YOLO-Pose 实时感知 → ShuffleNetV2+SA 跌倒分类 + 活动识别', indent=True)
add_para('2. 边缘端 Qwen2.5-1.5B/0.5B 语义增强 → 结构化事件 + 护理建议', indent=True)
add_para('3. TaskRouter 按置信度/复杂度/网络状态决策 → 复杂或低置信事件卸载云端', indent=True)
add_para('4. 云端 Qwen2.5-14B 二次研判 → confirm/reject/escalate 回传', indent=True)
add_para('5. 边缘按幂等规则写回状态 → 护士站展示（含 trace_id 全程关联）', indent=True)

# ===== 五、赛题指标对照 =====
add_heading('五、赛题硬指标达标情况（x86 实测基线）', 1)

make_table(
    ['指标', '竞赛要求', '边缘 1.5B', '边缘 0.5B', '云端 14B'],
    [Cm(2.5), Cm(3.5), Cm(3.2), Cm(3.2), Cm(3.4)],
    [
        ['国产大模型', '基于国产大模型', '✅ Qwen2.5', '✅ Qwen2.5', '✅ Qwen2.5'],
        ['单次推理内存', '≤ 1.5GB', '⚠️ 1658MB 超标', '✅ 516MB', '—（云端不限）'],
        ['TTFT', '减少 75%（<200ms）', '✅ 31.9ms', '✅ 19.8ms', '待测（vLLM）'],
        ['满血能力保持', '80-90%', '待蒸馏评测', '待蒸馏评测', '满血基准'],
        ['云边协同', '模型级协同', '✅ 轻量模型', '✅ 轻量模型', '✅ 全量模型'],
        ['蒸馏方案', '大→小压缩', '✅ 14B→1.5B', '✅ 14B→0.5B', '教师模型'],
        ['断网自治', '保持率 ≥90%', '✅ 离线决策', '✅ 离线决策', '—'],
    ],
)

doc.add_paragraph()
add_para('注：x86 数据仅作基线，Jetson 实机数据与蒸馏评测结果出来后更新本表；'
         '不得以 Windows/x86 结果描述为 Jetson 实测。', size=9, color=RGBColor(0x88, 0x88, 0x88))

# ===== 六、风险与备选 =====
add_heading('六、风险与备选', 1)

risks = [
    ('Jetson 内存超标', '1.5B 若在 Jetson 上峰值 RSS > 1.5GB，启用 0.5B compact 路径；保留 1.5B 质量对比数据，说明质量/资源权衡'),
    ('云端显存不足', '14B 显存不足时降级 Qwen2.5-7B 或 AWQ 4bit 量化；接口不变，仅换模型后端'),
    ('蒸馏效果不达标', '增大病房 NLU 语料、调整 KD 温度；蒸馏结果标为研究性增强，优先保障云边主链路'),
    ('真实模型接入延迟', 'vLLM 无法稳定启动时，继续用接口一致 mock consumer 打通闭环，视频/报告区分两种环境'),
]
for title, desc in risks:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{title}：')
    r.font.bold = True
    r.font.size = Pt(10)
    set_cjk(r)
    r = p.add_run(desc)
    r.font.size = Pt(10)
    set_cjk(r)

# 保存
out_path = r'C:\Users\Aarontang\Desktop\云边协同LLM模型选型方案-Qwen2.5.docx'
doc.save(out_path)
print('已生成：' + out_path)
