# -*- coding: utf-8 -*-
"""亚伦工作汇报 docx 生成器（简洁版）"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微软雅黑"
doc = Document()

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r.font.size = Pt(16)

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        r.font.size = Pt(13)

def p(text, bold=False, size=11, color=None, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_shading(cell, "1F4E79")
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r+1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()
p("智慧病房云边协同系统", bold=True, size=22, color=(0x1F,0x4E,0x79), align="center")
p("亚伦工作汇报", bold=True, size=18, color=(0x2E,0x75,0xB6), align="center")
doc.add_paragraph()
p("日期：2026 年 7 月", size=12, color=(0x59,0x59,0x59), align="center")
doc.add_page_break()

# ===== 一、任务概览 =====
h1("一、任务概览")
p("本人亚伦承担三项任务，整体完成情况如下：")
table(
    ["任务", "内容", "完成度", "状态"],
    [
        ["任务1", "边缘端识别模型选型 + 病房场景技术适配", "30%", "进行中"],
        ["任务2", "配合建鸿调研、具象化场景需求", "40%", "进行中"],
        ["任务3", "搭建项目整体技术框架 + 实时测试/决策/通信测试", "100%", "已完成"],
    ],
    [0.8, 4.0, 0.8, 1.0],
)
p("其中任务3（搭建整体技术框架）为核心成果，已完成端到端联调验证。任务1（模型选型）与任务2（场景具象化）部分完成，需后续推进。", size=10, color=(0x59,0x59,0x59))

# ===== 二、已完成工作 =====
h1("二、已完成工作")

h2("2.1 搭建项目整体技术框架（任务3，100%）")
p("从零搭建 smart-ward 独立项目，复用智能教室项目已验证的技术路线（Vue/FastAPI/MQTT/MySQL/SQLite/Docker Compose），将领域模型、数据契约、事件闭环全部面向智慧病房重新设计。")
table(
    ["模块", "内容", "规模"],
    [
        ["contracts/", "6 份 MQTT JSON Schema（消息契约联调标准）", "6 文件"],
        ["edge-agent/", "4 类采集适配器 + 融合引擎 + 推理空壳 + 场景驱动 + SQLite 缓存", "8 文件"],
        ["cloud-backend/", "11 张表 ORM + 21 个 REST API + WebSocket + MQTT 处理器", "7 文件"],
        ["training-coordinator/", "协同训练调度空壳（建鸿/振鑫接手）", "2 文件"],
        ["cloud-frontend/", "Vue 3 护士站工作台（3 列布局）", "4 文件"],
        ["docs/", "事件字典 + MQTT 契约 + 骨架交付说明 + 技术框架 docx", "5 文件"],
    ],
    [1.8, 4.0, 1.0],
)
p("代码规模约 4970 行，8 个 commit，17 项单元测试全绿。", bold=True)

h2("2.2 新增 10 项智能功能")
p("在框架基础上扩展 10 项病房智能功能，覆盖患者安全、护理响应、环境设备三类场景：")
table(
    ["类别", "功能", "事件类型", "优先级"],
    [
        ["患者安全", "坠床预警（事前预警）", "fall_prediction", "P1"],
        ["患者安全", "长时间静止检测", "long_still", "P2"],
        ["患者安全", "异常体态识别", "abnormal_posture", "P2"],
        ["患者安全", "抽搐检测", "seizure", "P1"],
        ["患者安全", "压疮预防提醒", "bedsore_risk", "P3"],
        ["患者安全", "设备故障预警", "device_fault", "P3"],
        ["护理响应", "交接班摘要自动生成", "-", "-"],
        ["环境设备", "环境自适应（夜间离床开夜灯）", "-", "-"],
        ["环境设备", "空气质量联动（CO₂ 超阈值开新风）", "-", "-"],
        ["环境设备", "床位占用可视化（含患者别名）", "-", "-"],
    ],
    [1.0, 2.6, 1.6, 0.8],
)

h2("2.3 端到端联调验证")
p("8 个 Docker 服务全部启动并通过联调验证：")
bullet("3 个边缘代理上报 observation + health，3 床在线")
bullet("边缘端自动触发 6 类新事件并入库（fall_prediction/long_still/abnormal_posture/seizure/bedsore_risk/device_fault）")
bullet("告警确认闭环：new -> notified -> acknowledged -> resolved，含处置记录与审计日志")
bullet("交接班摘要 API 生成 12 事件统计含分布")
bullet("环境控制 API 通过 MQTT 下发到边缘端")
bullet("前端 http://localhost:8081 可访问，3 列布局展示床位/告警/交接班")
p("联调中发现并修复 4 个真实 bug（环境变量覆盖、init.sql 未挂载、外键约束 flush、交接班时区），均已提交。", size=10, color=(0x59,0x59,0x59))

# ===== 三、技术成果 =====
h1("三、技术成果")

h2("3.1 复用策略")
p("验证了“非推倒重来”的复用策略，从智能教室项目复用基础设施，仅替换领域模型：")
table(
    ["方式", "数量", "说明"],
    [
        ["原样复制", "8 文件", "websocket_manager / logger / Dockerfile / 前端 proxy/nginx/vite 等"],
        ["模式复用+字段重写", "4 文件", "database / mqtt_handler / main / schemas（Room->Ward/Bed/EdgeNode）"],
        ["新写", "-", "11 表 init.sql / 新 MQTT 主题树 / 4 类适配器 / 融合引擎 / 6 份 Schema"],
    ],
    [1.5, 1.0, 4.0],
)

h2("3.2 数据模型")
p("云端 MySQL 共 11 张表，覆盖病房全场景：wards / beds / edge_nodes / observations / safety_events / alert_tasks / event_dispositions / model_versions / model_deployments / audit_logs / shift_summaries。")
p("事件状态机：new -> notified -> acknowledged -> resolved / false_positive / escalated，将识别事件、通知任务、人工处置分开建模。")

h2("3.3 MQTT 通信契约")
p("定义了完整的主题树与通用消息信封（message_id/event_id/schema_version/occurred_at/source/trace_id/payload），所有消息 QoS=1，消费端按 message_id 去重，边缘侧按序号补传。")
bullet("上行：ward/{ward}/node/{node}/{observation,event,health}")
bullet("下行：ward/{ward}/alert/{event_id}/ack + node/{node}/{config/set,model/deploy}")
bullet("训练隔离：training/{job_id}/node/{node}/{command,update}")

# ===== 四、未完成工作 =====
h1("四、未完成工作")

h2("4.1 边缘识别模型选型（任务1，30%）")
p("已就位：inference.py 预留模型接入点，强制输出方案书 §3.3 验收字段；4 类适配器跑通模拟数据。")
p("待完成：")
bullet("YOLO Nano 系 / 姿态模型 / OpenVINO vs TensorRT / INT8 量化对比文档")
bullet("真实模型接入 inference.py（当前是透传空壳）")
bullet("目标硬件实测延迟/FPS/召回率/误报率")

h2("4.2 配合建鸿调研（任务2，40%）")
p("已就位：6 份 JSON Schema 契约 + 事件字典 + 6 类场景跑通。")
p("待完成：")
bullet("完整项目调研报告（配合建鸿）")
bullet("细化病房画像场景需求文档")

# ===== 五、下一步计划 =====
h1("五、下一步计划")
table(
    ["时间", "工作", "对应任务"],
    [
        ["近期", "产出边缘模型选型对比文档（YOLOv8n-pose / OpenVINO / TensorRT）", "任务1"],
        ["近期", "在 inference.py 接入真实 YOLO 模型，目标硬件实测", "任务1"],
        ["中期", "配合建鸿完成完整项目调研报告", "任务2"],
        ["中期", "配合团队接入分布式协同训练与扩散模型", "团队协作"],
        ["8/31 前", "完成赛事申报材料技术部分", "建鸿协作"],
    ],
    [0.8, 4.2, 1.2],
)

# ===== 保存 =====
out = "docs/亚伦工作汇报.docx"
doc.save(out)
import os
print(f"✅ 已生成：{out} ({os.path.getsize(out)/1024:.1f} KB)")
