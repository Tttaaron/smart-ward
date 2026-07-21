# -*- coding: utf-8 -*-
"""硬件选型与云边端部署方案 docx 生成器"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微软雅黑"
doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r.font.size = Pt(16)

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        r.font.size = Pt(13)

def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        r.font.size = Pt(12)

def p(text, bold=False, size=11, color=None, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def bullet(text, bold_label=None):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    if bold_label:
        r1 = para.add_run(bold_label)
        r1.bold = True
        r1.font.name = FONT
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r1.font.size = Pt(11)
        r2 = para.add_run(text)
        r2.font.name = FONT
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r2.font.size = Pt(11)
    else:
        run = para.add_run(text)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(11)

def code(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_shading(cell, "1F4E79")
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r+1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()
p("智慧病房云边协同系统", bold=True, size=22, color=(0x1F,0x4E,0x79), align="center")
p("硬件选型与云边端部署方案", bold=True, size=18, color=(0x2E,0x75,0xB6), align="center")
doc.add_paragraph()
p("支持云端计算与端侧计算两种方案", size=13, color=(0x59,0x59,0x59), align="center")
doc.add_paragraph()
p("版本 V0.1", size=12, color=(0x40,0x40,0x40), align="center")
p("日期：2026 年 7 月", size=11, color=(0x59,0x59,0x59), align="center")
doc.add_page_break()

# ===== 第一章：两种计算方案概述 =====
h1("第一章 两种计算方案概述")

p("本项目支持云端计算与端侧计算两种方案，两种方案在数据流向、实时性、可靠性、部署成本上各有取舍。实际部署时可根据病房规模、网络条件、硬件预算灵活选择，或采用云边协同的混合方案。")

h2("1.1 方案对比")
table(
    ["维度", "方案 A：端侧计算", "方案 B：云端计算", "方案 C：云边协同（推荐）"],
    [
        ["推理位置", "边缘计算盒本地推理", "云端服务器推理", "边缘轻量 + 云端复核"],
        ["数据流向", "传感器→边缘盒→本地决策→MQTT上报", "传感器→边缘盒→云端→决策→下发", "边缘优先决策，云端按需复核"],
        ["实时性", "高（1秒内告警）", "中（依赖网络往返）", "高（边缘不等待云端）"],
        ["断网可用", "完全可用", "不可用", "可用（降级模式）"],
        ["模型规模", "轻量（YOLOv8n，6MB）", "大模型可（云端算力强）", "边缘轻量 + 云端复杂时序"],
        ["硬件成本", "中（每床位1个边缘盒）", "低（边缘只做采集）", "中高（边缘+云端都要算力）"],
        ["隐私保护", "好（视频不出本地）", "差（视频上传云端）", "好（仅传事件摘要）"],
        ["适合场景", "网络不稳、隐私敏感", "网络稳定、集中监控", "赛事展示、真实部署"],
    ],
    [1.0, 1.6, 1.6, 2.0],
)

h2("1.2 推荐方案：云边协同")
p("本项目采用方案 C（云边协同），核心原则：")
bullet("紧急事件（P1 跌倒/抽搐）边缘端本地决策，不等待云端返回，1 秒内告警。", "紧急优先本地：")
bullet("灰区事件（置信度低）上传脱敏片段供云端复核，云端返回建议优先级。", "灰区云端复核：")
bullet("断网时边缘端继续采集、识别、告警、本地缓存；恢复后补传事件，不丢不重。", "断网自治：")
bullet("原始视频留边缘端，仅上传事件摘要、脱敏截图；云端不长期保存完整视频。", "隐私默认：")

doc.add_page_break()

# ===== 第二章：硬件选型 =====
h1("第二章 硬件选型")

p("按功能反推硬件需求，而非泛泛采购。当前 10 项智能功能中，6 项依赖摄像头，3 项依赖床垫传感器，其余为可选补充。")

h2("2.1 按功能反推硬件")
table(
    ["功能", "必需硬件", "可选补充", "覆盖说明"],
    [
        ["跌倒检测", "RGB 摄像头", "毫米波雷达", "摄像头姿态识别为主"],
        ["坠床预警", "摄像头 + 床垫传感器", "-", "床沿位置 + 占床融合判定"],
        ["离床检测", "床垫压力传感器", "摄像头", "床垫是主源"],
        ["长时间静止", "摄像头", "床垫", "姿态变化跟踪"],
        ["异常体态", "摄像头", "-", "蜷缩/前倾/抓胸口姿态分类"],
        ["抽搐检测", "摄像头", "-", "关键点高频抖动"],
        ["压疮预防", "床垫压力传感器", "摄像头", "同一体位持续时长"],
        ["输液异常", "输液滴速传感器", "摄像头", "传感器为主，视觉补充"],
        ["呼叫按钮", "呼叫按钮", "-", "物理按键或无线按钮"],
        ["环境监测", "温湿度/CO₂/光照传感器", "-", "环境舒适度 + 联动控制"],
    ],
    [1.2, 2.2, 1.4, 2.4],
)

h2("2.2 硬件清单（按优先级）")

h3("2.2.1 必选（最小演示链路）")
table(
    ["硬件", "推荐型号", "价格", "用途", "接入方式"],
    [
        ["边缘计算盒", "Orange Pi 5（8GB）", "750-850 元", "运行 edge-agent + 模型推理", "Docker / Ethernet"],
        ["RGB 摄像头", "罗技 C920（1080p）", "350-400 元", "人体检测 + 姿态识别", "USB"],
        ["床垫压力传感器", "FSR402 + ESP32", "80-100 元", "占床/离床检测", "ESP32 MQTT 上报"],
    ],
    [1.3, 1.8, 1.0, 2.2, 1.5],
)

h3("2.2.2 推荐（完整功能）")
table(
    ["硬件", "推荐型号", "价格", "用途"],
    [
        ["呼叫按钮", "无线门铃改装 / ESP32 + 按钮", "50-200 元", "护士呼叫闭环"],
        ["环境传感器", "SCD30 CO₂ + BH1750 光照 + DHT22", "100-300 元", "环境监测 + 联动"],
        ["输液监测器", "滴速传感器 + ESP32", "200-500 元", "输液异常检测"],
    ],
    [1.3, 2.0, 1.0, 3.5],
)

h3("2.2.3 可选（赛事加分）")
table(
    ["硬件", "型号", "价格", "用途"],
    [
        ["毫米波雷达", "HLK-LD241", "200-500 元", "夜间/遮挡场景，隐私友好"],
        ["NPU 加速棒", "Coral USB Accelerator", "500-1000 元", "边缘推理加速"],
        ["PoE 交换机", "4 口 PoE", "200-400 元", "摄像头供电 + 数据合一"],
        ["UPS", "500VA 小型 UPS", "200-400 元", "断电保护"],
    ],
    [1.3, 2.0, 1.0, 3.5],
)

h2("2.3 边缘计算盒选型对比")
p("边缘盒是最关键决策，直接影响模型推理能力。以下是 4 种方案对比：")
table(
    ["方案", "硬件", "推理能力", "价格", "适合场景"],
    [
        ["A：纯 CPU", "树莓派 5 / Intel NUC", "YOLOv8n 约 5-10 FPS", "500-1500", "最小成本，规则+轻量模型"],
        ["B：NPU 加速（推荐）", "Orange Pi 5（RK3588）", "YOLOv8n 约 15-30 FPS", "750-850", "性价比最高，赛事推荐"],
        ["C：USB 加速棒", "树莓派 + Coral USB", "YOLOv8n 约 30+ FPS", "1000-1500", "需要高 FPS"],
        ["D：独立 GPU", "Jetson Nano / Orin Nano", "YOLOv8n 约 40+ FPS", "1500-3000", "赛事高规格"],
    ],
    [1.5, 1.8, 1.8, 0.8, 2.1],
)
p("推荐 Orange Pi 5：自带 6TOPS NPU，跑 YOLOv8n-pose INT8 量化后 15-30 FPS 够用；支持 Docker 直接部署当前框架；8GB 内存；OpenVINO/RKNN 都支持，模型选型灵活。", bold=True)

h2("2.4 1 床位最小演示预算")
table(
    ["项目", "型号", "单价", "数量", "小计"],
    [
        ["边缘计算盒", "Orange Pi 5 8GB", "800", "1", "800"],
        ["USB 摄像头", "罗技 C920", "400", "1", "400"],
        ["床垫传感器套件", "FSR402 + ESP32", "100", "1", "100"],
        ["SD 卡 32GB", "-", "30", "1", "30"],
        ["电源 + 网线 + 面包板", "-", "60", "1", "60"],
        ["合计", "", "", "", "约 1390 元"],
    ],
    [1.5, 1.8, 0.8, 0.8, 1.0],
)

h2("2.5 3 床位完整配置预算")
table(
    ["项目", "单价", "数量", "小计"],
    [
        ["Orange Pi 5 8GB", "800", "3", "2400"],
        ["USB 摄像头 C920", "400", "3", "1200"],
        ["床垫传感器套件", "100", "3", "300"],
        ["呼叫按钮 + ESP32", "150", "3", "450"],
        ["环境传感器套件", "250", "3", "750"],
        ["输液监测器", "400", "3", "1200"],
        ["云端服务器（用现有电脑）", "0", "1", "0"],
        ["PoE 交换机 + UPS", "600", "1", "600"],
        ["合计", "", "", "约 6900 元"],
    ],
    [2.0, 0.8, 0.8, 1.0],
)

doc.add_page_break()

# ===== 第三章：端侧部署方案 =====
h1("第三章 端侧部署方案")

p("端侧部署指在边缘计算盒（Orange Pi 5）上运行 edge-agent，完成本地采集、推理、决策、告警，并通过 MQTT 将事件摘要上报云端。")

h2("3.1 端侧部署架构")
code("┌─────────────────────────────────────────┐")
code("│      Orange Pi 5 边缘计算盒              │")
code("│                                          │")
code("│  ┌──────────┐  ┌──────────┐  ┌──────┐  │")
code("│  │ USB 摄像头│  │ ESP32 床垫│  │环境  │  │")
code("│  │ (OpenCV) │  │ (MQTT订阅)│  │传感器│  │")
code("│  └────┬─────┘  └────┬─────┘  └──┬───┘  │")
code("│       │             │           │       │")
code("│       ▼             ▼           ▼       │")
code("│  ┌─────────────────────────────────┐   │")
code("│  │      edge-agent 容器            │   │")
code("│  │  采集适配器 → 推理 → 融合 → 决策│   │")
code("│  │  → SQLite缓存 → MQTT上报       │   │")
code("│  └──────────────┬──────────────────┘   │")
code("│                 │                        │")
code("└─────────────────┼────────────────────────┘")
code("                  │ MQTT (QoS 1)")
code("                  ▼")
code("              云端事件中心")

h2("3.2 端侧部署步骤")

h3("3.2.1 Orange Pi 5 系统准备")
p("第 1 天完成。烧录 Ubuntu 22.04 系统，安装 Docker。")
code("# 1. 用 balenaEtcher 烧录 Ubuntu 22.04 到 SD 卡")
code("# 2. 首次启动，配置 Wi-Fi，SSH 登录")
code("# 3. 安装 Docker")
code("curl -fsSL https://get.docker.com | sh")
code("sudo usermod -aG docker $USER")
code("# 4. 验证")
code("docker --version")

h3("3.2.2 摄像头接入")
p("第 2 天完成。USB 摄像头即插即用，Linux 内核自带 UVC 驱动。")
code("# 验证摄像头识别")
code("ls /dev/video0   # 应存在")
code("v4l2-ctl --list-devices   # 列出摄像头设备")
code("")
code("# 安装 Python 依赖（用于后续推理）")
code("pip install opencv-python ultralytics")

h3("3.2.3 部署 edge-agent 容器")
p("第 3 天完成。拉取镜像或本地构建。")
code("# 克隆项目")
code("git clone <仓库地址> smart-ward")
code("cd smart-ward")
code("")
code("# 单独启动 edge-agent（云端在另一台机器）")
code("docker compose up -d edge-bed-01")
code("")
code("# 查看日志")
code("docker logs -f smart-ward-edge-bed-01-1")

h3("3.2.4 ESP32 床垫传感器接入")
p("第 3-4 天完成。ESP32 烧录固件，通过 Wi-Fi 连接到 Orange Pi 5 的 MQTT broker。")
bullet("ESP32 接线：3.3V → FSR402 一端 → 10kΩ 电阻分压 → GPIO34（ADC）", "接线：")
bullet("Arduino IDE 烧录固件，ESP32 连 Wi-Fi，每秒上报压力值到 MQTT topic sensor/bed/B01/pressure", "固件：")
bullet("edge-agent 的 BedSensorAdapter 订阅该 topic，解析压力值判断占床/离床", "订阅：")

h2("3.3 端侧配置（环境变量）")
p("Orange Pi 5 上的 docker-compose 配置，指向云端服务器 IP：")
code("services:")
code("  edge-bed-01:")
code("    build: ./edge-agent")
code("    environment:")
code("      WARD_ID: W-01")
code("      BED_ID: B01")
code("      EDGE_NODE_ID: EDGE-W01-B01")
code("      MQTT_BROKER: 192.168.1.100  # 云端服务器 IP")
code("      MQTT_PORT: 1884")
code("      MODEL_NAME: yolov8n-pose")
code("      MODEL_VERSION: 1.0.0-int8")
code("      SCENARIO_PROFILE: ''  # 真实硬件模式，关闭场景模拟")
code("    devices:")
code("      - /dev/video0:/dev/video0  # 挂载摄像头")
code("    volumes:")
code("      - edge-b01-data:/app/data")

h2("3.4 端侧断网自治")
p("端侧部署的核心优势：网络中断时仍可工作。")
bullet("摄像头/传感器继续采集，edge-agent 继续推理与告警", "采集不停：")
bullet("事件写入本地 SQLite，标记 synced=0", "缓存不丢：")
bullet("MQTT 恢复后按序号补传，云端按 event_id 去重", "补传不重：")
bullet("断网期间告警通过本地护士站（局域网）或蜂鸣器触发", "告警可达：")

doc.add_page_break()

# ===== 第四章：云端部署方案 =====
h1("第四章 云端部署方案")

p("云端部署指 cloud-backend + cloud-frontend + mqtt-broker + mysql 运行在一台服务器上，接收所有边缘节点上报的事件，提供护士站工作台、事件查询、告警确认、模型下发等功能。")

h2("4.1 云端部署架构")
code("┌─────────────────────────────────────────────┐")
code("│          云端服务器（你的电脑 / 服务器）        │")
code("│                                              │")
code("│  ┌──────────┐  ┌──────────┐  ┌──────────┐  │")
code("│  │ Mosquitto│  │  MySQL   │  │ FastAPI  │  │")
code("│  │MQTT Broker│ │ 数据库   │  │事件中心  │  │")
code("│  │  :1884   │  │  :3308   │  │  :8001   │  │")
code("│  └────┬─────┘  └────┬─────┘  └────┬─────┘  │")
code("│       │             │             │          │")
code("│       └─────────────┼─────────────┘          │")
code("│                     │                        │")
code("│              ┌──────┴──────┐                 │")
code("│              │ Vue 护士站  │                 │")
code("│              │  :8081      │                 │")
code("│              └─────────────┘                 │")
code("└──────────────────────────────────────────────┘")
code("         ▲ MQTT 上报         ▲ WebSocket 推送")
code("         │                   │")
code("    边缘节点 1-3         护士站浏览器")

h2("4.2 云端部署步骤")

h3("4.2.1 启动云端服务")
p("在云端服务器（你的电脑即可）执行：")
code("cd smart-ward")
code("docker compose up -d mqtt-broker mysql cloud-backend cloud-frontend")
code("")
code("# 验证服务")
code("curl http://localhost:8001/health   # 应返回 {\"status\":\"ok\"}")
code("curl http://localhost:8001/api/stats  # 返回统计")

h3("4.2.2 访问护士站")
p("浏览器打开：")
bullet("护士站工作台：http://localhost:8081")
bullet("API 文档：http://localhost:8001/docs")
bullet("MQTT 端口：localhost:1884")

h3("4.2.3 云端计算模式（可选）")
p("若采用方案 B（云端计算），边缘端仅做采集，视频流上传云端推理：")
bullet("边缘 edge-agent 采集摄像头帧，通过 MQTT 或 RTSP 上传云端", "采集上传：")
bullet("云端 cloud-backend 接收视频流，调用 YOLO 模型推理", "云端推理：")
bullet("云端生成事件，通过 WebSocket 推送到护士站", "云端决策：")
bullet("边缘端只负责采集与执行云端下发的控制指令", "边缘执行：")
p("注意：云端计算模式对网络带宽要求高（视频流），且断网时完全不可用。推荐仅在方案 C 云边协同中用于灰区事件复核。", size=10, color=(0x59,0x59,0x59))

h2("4.3 云端服务器要求")
table(
    ["配置项", "最低要求", "推荐配置", "说明"],
    [
        ["CPU", "4 核", "8 核+", "云端推理时需更强"],
        ["内存", "8 GB", "16 GB", "MySQL + 后端 + 模型"],
        ["硬盘", "50 GB", "100 GB SSD", "事件数据 + 日志"],
        ["网络", "100 Mbps", "千兆有线", "多路视频流需高带宽"],
        ["GPU（可选）", "无", "NVIDIA GTX 1660+", "云端大模型推理加速"],
    ],
    [1.2, 1.3, 1.5, 2.5],
)
p("当前演示阶段：你的开发电脑即可作为云端服务器，无需额外采购。")

doc.add_page_break()

# ===== 第五章：端到云端数据传输 =====
h1("第五章 端到云端数据传输")

h2("5.1 数据流向总览")
code("传感器层          边缘层              云端层            展示层")
code("┌────────┐    ┌────────────┐    ┌────────────┐    ┌──────┐")
code("│ 摄像头  │───→│            │    │            │    │      │")
code("│ 床垫    │───→│  edge-agent│───→│ MQTT Broker│───→│ FastAPI")
code("│ 输液    │───→│            │MQTT│            │WS  │      │")
code("│ 环境    │───→│  推理+融合  │QoS1│            │    │  ↓   ")
code("│ 呼叫    │───→│            │    │            │    │ Vue  │")
code("└────────┘    └─────┬──────┘    └─────┬──────┘    │ 护士站│")
code("                    │                  │           └──────┘")
code("                    │ SQLite           │ MySQL")
code("                    │ 本地缓存          │ 持久化")
code("                    └──────────────────┘")

h2("5.2 MQTT 主题与传输协议")
p("所有端到云端通信走 MQTT，QoS=1 保证至少一次送达。")
table(
    ["方向", "主题", "内容", "频率", "大小"],
    [
        ["上行", "ward/{ward}/node/{node}/observation", "多源观测数据", "每 3 秒", "约 2 KB"],
        ["上行", "ward/{ward}/node/{node}/event", "安全事件（跌倒/离床等）", "事件触发", "约 1 KB"],
        ["上行", "ward/{ward}/node/{node}/health", "节点健康心跳", "每 30 秒", "约 0.5 KB"],
        ["下行", "ward/{ward}/alert/{event_id}/ack", "告警确认指令", "护士操作", "约 0.5 KB"],
        ["下行", "node/{node}/config/set", "环境控制指令", "联动触发", "约 0.3 KB"],
        ["下行", "node/{node}/model/deploy", "模型版本下发", "版本更新", "约 1 KB"],
    ],
    [0.8, 2.5, 2.0, 1.0, 0.8],
)

h2("5.3 消息信封（通用格式）")
p("所有消息外层统一信封，含追踪与去重字段：")
code("{")
code("  \"message_id\": \"550e8400-e29b-41d4-a716-446655440000\",")
code("  \"event_id\": \"660e8400-...\",     // 关联事件ID，非事件为null")
code("  \"schema_version\": \"v1\",")
code("  \"occurred_at\": \"2026-07-21T08:30:00Z\",")
code("  \"source\": \"edge:EDGE-W01-B01\",")
code("  \"trace_id\": \"770e8400-...\",")
code("  \"payload\": { ... }              // 业务数据")
code("}")

h2("5.4 传输可靠性保障")
table(
    ["机制", "说明", "实现位置"],
    [
        ["QoS 1", "至少一次送达，消费端按 message_id 去重", "MQTT Broker + 消费端"],
        ["断网缓存", "事件写本地 SQLite，标记 synced=0", "edge-agent LocalDatabase"],
        ["恢复补传", "MQTT 重连后按序号补传未同步事件", "edge-agent sync_offline_data()"],
        ["幂等入库", "云端按 event_id 唯一约束去重", "cloud-backend mqtt_handler"],
        ["遗嘱消息", "边缘节点断开时 Broker 自动发布 offline", "MQTT LWT（后续实现）"],
    ],
    [1.2, 3.3, 2.5],
)

h2("5.5 隐私保护与数据脱敏")
bullet("原始视频留边缘端，默认不长期保存，仅事件触发时截取脱敏片段", "视频不出本地：")
bullet("事件上传仅含 event_type/confidence/rule_hits，不含人脸图像", "只传摘要：")
bullet("云端展示用床位号 + 匿名别名（张阿姨/李伯伯），不存真实姓名", "匿名别名：")
bullet("前端默认开启人脸模糊，可配置隐私遮挡区", "人脸模糊：")
bullet("误报标记回流训练前必须脱敏 + 人工审核", "数据回流脱敏：")

doc.add_page_break()

# ===== 第六章：云边协同混合方案 =====
h1("第六章 云边协同混合方案")

p("实际部署推荐云边协同：边缘端做实时识别与告警，云端做事件复核、模型训练与下发、数据分析。两种计算不是二选一，而是分工协作。")

h2("6.1 协同决策流程")
code("边缘端检测到事件")
code("       │")
code("       ▼")
code("┌───────────────────┐")
code("│ 置信度 ≥ 紧急阈值? │")
code("└───────┬───────────┘")
code("        │")
code("   是 ──┤── 否")
code("        │      │")
code("        ▼      ▼")
code("  本地立即告警   上传脱敏片段")
code("  上报事件摘要   云端复杂模型复核")
code("  (不等云端)     │")
code("        │      ▼")
code("        │   返回建议优先级")
code("        │      │")
code("        └──┬───┘")
code("           ▼")
code("      护士站展示")
code("      (本地+复核结果)")

h2("6.2 分层职责")
table(
    ["层级", "职责", "算力要求", "数据量"],
    [
        ["边缘端", "实时采集 + 轻量推理 + 即时告警 + 本地缓存", "低（NPU 6TOPS）", "原始视频本地处理"],
        ["云端", "复杂复核 + 模型训练 + 版本下发 + 数据分析 + 护士站", "中高（CPU/GPU）", "仅事件摘要上传"],
        ["展示端", "护士站工作台 + 交接班 + 告警确认", "低（浏览器）", "WebSocket 增量推送"],
    ],
    [1.0, 3.0, 1.5, 1.5],
)

h2("6.3 模型版本管理")
p("云端训练新模型后，通过 MQTT 下发到边缘端，支持灰度发布与回滚：")
bullet("云端训练 → 验证集评估 → 标记 released", "训练：")
bullet("通过 model/deploy 主题下发到指定节点（灰度）", "下发：")
bullet("边缘端下载制品、校验 checksum、加载新模型", "加载：")
bullet("异常时一键回滚到上一版本", "回滚：")
bullet("每次部署记录到 model_deployments 表", "留痕：")

doc.add_page_break()

# ===== 第七章：部署时间规划 =====
h1("第七章 部署时间规划")

h2("7.1 1 床位最小演示（5 天落地）")
table(
    ["天数", "任务", "产出", "负责"],
    [
        ["Day 1", "硬件到货 + Orange Pi 5 烧录 Ubuntu + 装 Docker", "设备就绪", "亚伦"],
        ["Day 2", "USB 摄像头接入 + OpenCV 读帧验证", "摄像头帧可见", "亚伦"],
        ["Day 3", "YOLOv8n-pose 模型推理接入 + posture 输出", "真实姿态识别", "亚伦"],
        ["Day 4", "ESP32 床垫传感器接入 + MQTT 上报", "占床/离床真实数据", "亚伦"],
        ["Day 5", "端云联调 + 前端验证 + 录制演示", "1 床位完整演示", "亚伦+景彬"],
    ],
    [0.8, 3.5, 2.0, 0.8],
)

h2("7.2 3 床位完整配置（10 天）")
table(
    ["阶段", "天数", "任务", "产出"],
    [
        ["阶段 1", "Day 1-5", "1 床位最小演示（见上）", "1 床位跑通"],
        ["阶段 2", "Day 6-8", "扩展到 3 床位 + 呼叫按钮 + 环境传感器", "3 床位硬件就位"],
        ["阶段 3", "Day 9-10", "全功能联调 + 压力测试 + 演示彩排", "完整演示就绪"],
    ],
    [1.0, 1.0, 3.5, 1.5],
)

h2("7.3 到货前可提前准备")
bullet("下载 Ubuntu 22.04 镜像 + balenaEtcher 烧录工具", "系统镜像：")
bullet("pip install ultralytics，预下载 yolov8n-pose.pt 模型（7MB）", "模型文件：")
bullet("安装 Arduino IDE + ESP32 开发环境，熟悉 ESP32 烧录流程", "ESP32 环境：")
bullet("git clone 项目到云端服务器，docker compose up 验证云端服务", "云端验证：")

doc.add_page_break()

# ===== 第八章：验收标准与风险 =====
h1("第八章 验收标准与风险")

h2("8.1 端侧部署验收标准")
table(
    ["验收项", "验证方法", "通过标准"],
    [
        ["Docker 运行", "docker ps", "edge-agent 容器 Up"],
        ["摄像头识别", "ls /dev/video0", "设备存在"],
        ["摄像头读帧", "OpenCV VideoCapture(0)", "返回非空帧"],
        ["模型推理", "YOLOv8n-pose 推理一帧", "输出 posture + fall_score"],
        ["ESP32 上报", "mosquitto_sub 订阅 sensor/bed/B01/pressure", "每秒收到压力值"],
        ["断网缓存", "断开网络后触发事件，恢复后查 SQLite", "synced=0 事件补传为 synced=1"],
    ],
    [1.5, 3.0, 2.0],
)

h2("8.2 云端部署验收标准")
table(
    ["验收项", "验证方法", "通过标准"],
    [
        ["服务启动", "docker compose ps", "4 服务全 Up"],
        ["API 健康", "curl /api/health", "{\"status\":\"ok\"}"],
        ["事件入库", "curl /api/events", "返回事件列表"],
        ["前端访问", "浏览器 http://localhost:8081", "护士站页面显示"],
        ["WebSocket", "前端控制台", "收到 ping/pong 心跳"],
    ],
    [1.5, 3.0, 2.0],
)

h2("8.3 风险与应对")
table(
    ["风险", "触发条件", "应对", "负责人"],
    [
        ["Orange Pi 5 NPU 驱动装不上", "RKNN 工具链兼容问题", "退而求其次用 CPU 推理（5-10 FPS 够演示）", "亚伦"],
        ["USB 摄像头夜间画质差", "低照度场景", "赛事演示白天进行；后续加补光灯/毫米波雷达", "亚伦"],
        ["压力传感器阈值难调", "个体差异大", "先用串口打印 raw 值观察，再定阈值", "亚伦"],
        ["ESP32 Wi-Fi 不稳", "信号弱", "用网线连接；或换 ESP32-CAM", "亚伦"],
        ["云端带宽不足", "多路视频上传", "采用云边协同，仅传事件摘要不传视频", "亚伦"],
        ["断网演示失败", "网络环境不可控", "本地护士站走局域网；录屏备份", "景彬"],
    ],
    [1.8, 1.8, 2.5, 0.8],
)

doc.add_page_break()

# ===== 第九章：总结 =====
h1("第九章 总结")

h2("9.1 方案核心")
bullet("采用云边协同混合方案：边缘端做实时识别与告警，云端做复核、训练与下发", "架构：")
bullet("Orange Pi 5（RK3588 NPU）+ USB 摄像头 + ESP32 床垫传感器，1 床位约 1400 元", "硬件：")
bullet("MQTT QoS 1 + SQLite 断网缓存 + event_id 幂等去重，保证不丢不重", "传输：")
bullet("原始视频留边缘端，仅传事件摘要；前端匿名别名；人脸模糊默认开启", "隐私：")

h2("9.2 两种方案适用场景")
table(
    ["场景", "推荐方案", "理由"],
    [
        ["赛事演示", "云边协同（方案 C）", "体现边缘自治 + 云端协同创新点"],
        ["网络稳定医院", "云端计算（方案 B）", "集中管理，硬件成本低"],
        ["网络不稳/隐私敏感", "端侧计算（方案 A）", "断网可用，视频不出本地"],
        ["真实部署", "云边协同（方案 C）", "兼顾实时性、可靠性、隐私"],
    ],
    [1.5, 2.0, 3.0],
)

h2("9.3 下一步")
bullet("采购 Orange Pi 5 + USB 摄像头 + ESP32 传感器套件（约 1400 元，1 周到货）", "立即：")
bullet("到货后按本方案部署，5 天跑通 1 床位真实数据闭环", "近期：")
bullet("扩展到 3 床位 + 完善云端模型管理 + 协同训练", "中期：")
bullet("接入真实病房（需院方授权 + 网络隔离 + 合规审查）", "远期：")

# ===== 保存 =====
out = "docs/硬件选型与云边端部署方案.docx"
doc.save(out)
import os
print(f"✅ 已生成：{out} ({os.path.getsize(out)/1024:.1f} KB)")
