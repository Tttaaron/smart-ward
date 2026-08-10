# -*- coding: utf-8 -*-
"""亚伦本周工作汇报 docx 生成器（2026-07-28 至 2026-08-03）"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微软雅黑"
doc = Document()

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r.font.size = Pt(16)

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        r.font.size = Pt(13)

def p(text, bold=False, size=11, color=None, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_shading(cell, "1F4E79")
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r+1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()
p("智慧病房云边协同系统", bold=True, size=22, color=(0x1F,0x4E,0x79), align="center")
p("亚伦本周工作汇报", bold=True, size=18, color=(0x2E,0x75,0xB6), align="center")
doc.add_paragraph()
p("汇报周期：2026-07-28 至 2026-08-03", size=12, color=(0x59,0x59,0x59), align="center")
p("汇报人：唐亚伦（P1 边缘AI + 框架）", size=12, color=(0x59,0x59,0x59), align="center")
doc.add_page_break()

# ===== 一、本周概览 =====
h1("一、本周概览")
p("本周（7/28-8/3）围绕 T1 边缘 LLM 与 T2 协同推理两大主线，完成真实 YOLO 视觉管线、日常活动识别、LLM 友好日志精简与云边协同端到端联调。重点突破是定位并修复云端推理响应 wait_for_publish 自锁 bug（响应延迟 5s → 60ms），打通“边缘卸载 → 云端研判 → 回传 → 状态写回”完整闭环，并将全部成果合并入 master。")
table(
    ["指标", "数值"],
    [
        ["提交数", "8 个（含合并 master）"],
        ["edge-agent 单元测试", "从 26 项增至 63 项全绿"],
        ["云边协同响应延迟", "5.0s → 0.060s"],
        ["日志压缩比", "97%（128KB → 2.7KB）"],
        ["活动识别", "6 类（玩手机/吃饭/行走/睡眠/静坐/站立）"],
    ],
    [2.5, 4.0],
)

# ===== 二、本周核心工作 =====
h1("二、本周核心工作")

h2("2.1 真实 YOLO/YOLO-Pose 视觉管线（边缘感知升级）")
bullet("新增 YoloCameraAdapter（CAMERA_MODE=yolo 切换真实摄像头）、IoUTracker（无依赖 IoU 人员跟踪）、BehaviorAnalyzer（连续帧姿态/跌倒分/抖动分/静止时长）。")
bullet("新增 yolo_realtime_viewer.py 实时可视化窗口：检测框标注、行为面板、会话日志（帧节流/行为变化/跌倒事件/暂停/退出摘要），已用 Iriun/DroidCam 实机验证。")
bullet("Dockerfile 支持 INSTALL_YOLO 构建参数，可选依赖 requirements-yolo.txt 分离，默认镜像保持轻量。")

h2("2.2 日常活动识别（关键点规则分类，创新点）")
bullet("新增 activity_tracker.py：识别玩手机 / 吃饭 / 行走 / 睡眠 / 静坐 / 站立 6 类活动。")
bullet("复用 BehaviorAnalyzer 平滑 posture + 墙钟时长累积，修复“站立/静坐永远识别不出”的问题。")
bullet("连续 5 帧滞回去抖：实测真实摄像头日志 7 次输入翻转 → 仅 2 次合理输出切换，单帧闪现与短噪声串被吸收。")

h2("2.3 LLM 友好日志精简器")
bullet("新增 summarize_yolo_log.py：500 行帧级日志 → 46 行事件摘要（128KB → 2.7KB，压缩 97%）。")
bullet("输出会话统计、活动时长占比、活动时间段、跌倒事件，可直接喂给 LLM 生成护理汇报。")

h2("2.4 云边协同端到端联调（本周最大成果）")
bullet("打通完整闭环：边缘 TaskRouter 卸载低置信度事件 → MQTT → 云端研判（confirm/reject/escalate）→ 回传 → 边缘状态写回。")
bullet("定位并修复 wait_for_publish 自锁 bug：云端响应固定延迟 5 秒 → 60ms（根因：MQTT 回调线程内同步等待 PUBACK 造成自锁）。")
bullet("超时参数调优（2s → 5s）、Docker 构建修复（requirements-llm 拆分、PYTHONUNBUFFERED）、端口冲突处理。")

h2("2.5 版本管理与文档")
bullet("完成 YOLO 管线 / 活动识别 / 联调修复等提交，合并入 master 并推送 GitHub。")
bullet("文档测试数字统一为实测值（edge 63 + training 12）；修复测试夜间时间敏感问题（22 点后 bed_leave 被升级为 night_wandering 导致失败）。")

# ===== 三、关键问题解决 =====
h1("三、关键问题解决")
table(
    ["问题", "根因", "解决"],
    [
        ["云端推理响应固定延迟 5s", "wait_for_publish 在 MQTT 回调线程内自锁（PUBACK 由同一 loop 线程处理）", "移除同步等待，响应延迟降至 60ms"],
        ["站立/静坐识别不出", "ActivityRecognizer 重算姿态抖动 + 历史窗口仅 0.8s", "复用平滑 posture + 墙钟时长累积"],
        ["活动识别单帧闪现", "raw 标签在 phone/standing 间交替，无法连续 3 帧", "连续 5 帧滞回确认"],
        ["Docker 构建失败", "llama-cpp-python 在 slim 镜像需源码编译", "拆分 requirements-llm，INSTALL_LLM 才装"],
        ["夜间测试失败", "22 点后 _is_night() 为真，bed_leave 升级 night_wandering", "测试固定非夜间时段"],
    ],
    [2.2, 2.6, 2.2],
)

# ===== 四、当前任务完成度 =====
h1("四、当前任务完成度")
table(
    ["任务", "内容", "完成度", "状态"],
    [
        ["T1 边缘 LLM", "双路径（1.5B/0.5B）+ 语义增强/离线决策", "90%", "🟡 Jetson 实测待做"],
        ["T2 协同推理（边缘侧）", "路由/请求关联/超时回退/幂等", "90%", "🟡 云端 vLLM 待接"],
        ["真实视觉管线", "YOLO + 跟踪 + 行为分析", "100%", "✅ 完成"],
        ["活动识别", "6 类日常活动 + 去抖", "90%", "🟡 待接入 MQTT 上报"],
        ["日志精简", "LLM 友好会话摘要", "100%", "✅ 完成"],
    ],
    [1.4, 3.0, 0.9, 1.7],
)

# ===== 五、下周计划 =====
h1("五、下周计划（8/4-8/10）")
bullet("Jetson Orin Nano 实机部署与性能实测（TTFT / 内存 / 与 YOLO 并行资源占用）")
bullet("对比实验框架：纯边缘 vs 云边协同 vs 纯云端（同一数据集）")
bullet("断网业务保持率测试（目标 ≥ 90%）")
bullet("活动识别接入 MQTT 上报 + LLM 活动汇报闭环")
bullet("配合 P5 先伟完成病房 NLU 数据集与三路线评测")

# ===== 六、风险与需要支持 =====
h1("六、风险与需要支持")
table(
    ["风险", "影响", "需要支持"],
    [
        ["Jetson 设备未到位", "实机性能数据无法产出", "确认 Jetson 到货时间"],
        ["云端 14B 无 GPU", "真实 vLLM 链路无法验证", "提供 GPU 机器或先以 mock 联调"],
        ["训练侧滞后", "模型更新闭环演示缺失", "振鑫推进 FedAvg + 蒸馏"],
    ],
    [2.0, 2.5, 2.5],
)

doc.save("亚伦本周工作汇报.docx")
print("✓ 亚伦本周工作汇报.docx 生成成功")
