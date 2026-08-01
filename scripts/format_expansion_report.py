# -*- coding: utf-8 -*-
"""优化排版：拓展改进报告.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 默认样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cjk(run, name='微软雅黑'):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_cjk(r)
        if level == 1:
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
        elif level == 2:
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(0x33, 0x55, 0x99)
        elif level == 3:
            r.font.size = Pt(11.5)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return h

def add_para(text, bold=False, color=None, indent=True, size=11, italic=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    set_cjk(r)
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_cjk(r)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x0d, 0x6e, 0x5f)
    set_cjk(r)
    return p

def set_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def make_table(headers, rows, widths, header_color='1A56C4'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text, w in zip(t.rows[0].cells, headers, widths):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cjk(r)
        set_shading(cell, header_color)
    for row_data in rows:
        row = t.add_row()
        for cell, text, w in zip(row.cells, row_data, widths):
            cell.width = w
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(9)
            set_cjk(r)
    doc.add_paragraph()
    return t

# ===== 封面 =====
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('智慧病房云边协同系统')
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
set_cjk(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('功能拓展改进报告')
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x33, 0x55, 0x99)
set_cjk(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(15)
r = p.add_run('-- 基于视觉AI的患者行为理解与智能护理拓展')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
set_cjk(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('拓展方向：日常活动识别 / 睡眠质量评估 / 医护行为分析\n技术路线：YOLOv8-pose + Qwen2.5-1.5B + 云边协同')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
set_cjk(r)

doc.add_page_break()

# ===== 一、改进背景与目标 =====
add_heading('一、改进背景与目标', 1)

add_heading('1.1 现有系统局限', 2)
add_para('当前系统以"事件告警"为核心，只在发生安全事件（跌倒/离床/抽搐等）时触发响应。存在以下不足：')
add_bullet('缺乏持续性感知：系统只在"出事"时工作，无法反映患者日常状态')
add_bullet('缺乏语义理解：输出仅为结构化事件，无自然语言描述与护理建议')
add_bullet('缺乏时序分析：无法从长时间行为模式中发现趋势性问题')
add_bullet('缺乏医护行为记录：护士巡视/操作无自动记录，依赖人工')

add_heading('1.2 改进目标', 2)
add_para('将系统从"被动告警器"升级为"主动感知的AI护理助手"，实现：')
add_bullet('持续性行为理解：24小时不间断识别患者日常活动，主动生成护理摘要')
add_bullet('智能睡眠监护：基于姿态序列评估睡眠质量，辅助临床决策')
add_bullet('医护行为自动记录：自动统计巡视频率、操作时长，减轻护士记录负担')

add_heading('1.3 技术基础', 2)
add_para('本次拓展完全基于现有技术栈，无需额外硬件：')
add_bullet('数据源：病房摄像头（已有）')
add_bullet('姿态推理：YOLOv8-pose（已集成）')
add_bullet('语义理解：Qwen2.5-1.5B-INT4（已集成）')
add_bullet('云边协同：TaskRouter + MQTT（已实现）')

# ===== 二、功能一：日常活动识别 =====
add_heading('二、功能一：患者日常活动识别与LLM汇报', 1)

add_heading('2.1 功能概述', 2)
add_para('通过摄像头持续采集患者姿态序列，利用滑动窗口分析推断日常活动类型，并由边缘端轻量LLM生成自然语言汇报。')

add_heading('2.2 活动类型定义', 2)
make_table(
    ['活动类型', '姿态特征', '判定条件', '优先级'],
    [
        ['睡眠/休息', 'lying + 低体动', '卧床姿态持续>5分钟，体动<2次/分钟', '低'],
        ['静坐', 'sitting + 低位移', '坐位持续>2分钟，bbox位移<0.05', '低'],
        ['进食', 'sitting + 手部活动', '坐位 + 手部关键点在面部区域频繁', '中'],
        ['行走/活动', 'standing + 位移', 'bbox中心点连续帧位移>0.1', '中'],
        ['如厕', '站立->消失->站立', '人物bbox消失1-5分钟后恢复', '中'],
        ['康复训练', '重复性动作', '关键点角度周期性变化', '中'],
        ['跌倒', 'falling', 'fall_score>0.5 + 姿态突变', '紧急'],
    ],
    [Cm(2.5), Cm(3), Cm(6), Cm(2)]
)

add_heading('2.3 LLM汇报模式', 2)

add_heading('模式A：实时播报（活动切换时触发）', 3)
add_para('触发条件：活动类型发生变化时，边缘LLM生成一句话描述。', size=10)
add_code('Prompt: 患者B01在14:32从[睡眠]切换为[坐起]，已持续卧床95分钟。')
add_code('LLM输出: B01患者于14:32从睡眠中醒来坐起，建议观察是否需协助下床活动。')

add_heading('模式B：时段摘要（每30分钟触发）', 3)
add_para('边缘LLM汇总最近30分钟活动记录，生成结构化摘要。', size=10)
add_code('Prompt: 患者B01最近30分钟活动: 睡眠25分钟,坐起5分钟。请生成护理摘要。')
add_code('LLM输出: 14:00-14:30活动摘要：患者主要处于睡眠状态，未见异常。')

add_heading('模式C：交班报告（班次结束时触发）', 3)
add_para('云端14B模型汇总全班活动数据，生成可直接朗读的交班报告。', size=10)
add_code('LLM输出: B01白班总览：睡眠6.2h，进食2次，下床活动3次，康复训练完成度80%，无异常事件。')

add_heading('2.4 数据流与架构', 2)
add_code('摄像头 -> YOLOv8-pose(姿态检测) -> ActivityTracker(滑动窗口分类)')
add_code('-> LLMAdvisor(语义增强) -> MQTT上报 -> 云端存储 -> 前端活动日志面板')

add_heading('2.5 实现方案', 2)
add_para('新增模块：edge-agent/src/activity_tracker.py')
add_bullet('ActivityTracker类：维护最近N帧姿态记录，滑动窗口分类活动')
add_bullet('活动切换检测：当前后5帧姿态一致且与上一活动不同时触发切换')
add_bullet('定时摘要：每30分钟调用LLM生成时段报告')
add_bullet('集成到main.py主循环，每帧更新，不影响现有告警流程')
add_para('工时估算：约1.5天（核心逻辑200行 + 集成调试）', bold=True, color=RGBColor(0xe6, 0x7e, 0x22))

# ===== 三、功能二：智能睡眠质量评估 =====
add_heading('三、功能二：智能睡眠质量评估', 1)

add_heading('3.1 功能概述', 2)
add_para('基于姿态序列的长时间统计，评估患者睡眠质量，由LLM生成睡眠报告。')

add_heading('3.2 评估指标', 2)
make_table(
    ['指标', '计算方法', '临床意义'],
    [
        ['入睡时间', '灯灭后首次连续卧床>10分钟的时刻', '评估入睡困难'],
        ['睡眠时长', '卧床态势累计时间', '总体睡眠量'],
        ['翻身次数', '姿态变化次数（lying左右侧切换）', '睡眠深度指标'],
        ['夜间清醒次数', '坐起/站立次数（非入厕类）', '睡眠破碎程度'],
        ['晨醒时间', '首次持续坐起或站立的时刻', '作息规律性'],
        ['睡眠质量评分', '加权综合得分(0-100)', '总体评估'],
    ],
    [Cm(3), Cm(6), Cm(4.5)]
)

add_heading('3.3 LLM睡眠报告示例', 2)
add_code('输入: 患者B01昨夜睡眠数据: 入睡22:15, 睡眠时长6h45min, 翻身12次, 夜间清醒3次, 晨醒05:00')
add_code('LLM输出: B01昨夜睡眠质量评分72/100。入睡时间正常，但夜间清醒3次、翻身频繁，')
add_code('         提示睡眠较浅。建议关注患者是否有疼痛或不适，必要时调整用药。')

add_heading('3.4 实现方案', 2)
add_bullet('SleepTracker类：维护夜间全部姿态记录，计算各项指标')
add_bullet('每日晨醒后自动调用LLM生成睡眠报告')
add_bullet('报告存储到云端数据库，前端展示趋势图')
add_para('工时估算：约1天（与活动识别共用ActivityTracker基础架构）', bold=True, color=RGBColor(0xe6, 0x7e, 0x22))

# ===== 四、功能三：医护/访客行为分析 =====
add_heading('四、功能三：医护/访客行为分析', 1)

add_heading('4.1 功能概述', 2)
add_para('通过摄像头检测进入病房的人员，区分护士与访客，自动记录巡视频率、操作时长、访客时段等。')

add_heading('4.2 检测项目', 2)
make_table(
    ['检测项', '技术实现', '输出示例'],
    [
        ['护士巡视记录', '检测护士服装人员进入/离开', '14:20 护士巡视B01，停留8分钟'],
        ['巡视频率统计', '统计每小时进入次数', '过去2小时巡视3次，间隔正常'],
        ['护理操作识别', '检测输液/换药/翻身动作', '14:25 输液操作，时长12分钟'],
        ['访客管理', '非护士服装人员进出记录', '15:00 访客1人进入，停留15分钟'],
        ['夜间无护理预警', '夜间无人员进入检测', '夜间无护理人员进入，建议关注'],
        ['配护人员检测', '夜间非护士人员在床边', '夜间有配护人员在床边看护'],
    ],
    [Cm(3), Cm(5), Cm(5.5)]
)

add_heading('4.3 LLM汇报示例', 2)
add_code('Prompt: B01过去2小时人员进出: 护士3次(平均停留8分钟), 访客1次(15分钟), 无护理操作。')
add_code('LLM输出: B01过去2小时巡视3次，间隔正常。有1名访客探望。')
add_code('         未观察到输液/换药等护理操作。')

add_heading('4.4 实现方案', 2)
add_bullet('StaffTracker类：维护进出记录，区分护士/访客（服装颜色/人数）')
add_bullet('每小时生成一次汇报，存储到云端')
add_bullet('夜间特殊逻辑：无人员进入时触发预警')
add_para('工时估算：约1天（与活动识别共用基础架构）', bold=True, color=RGBColor(0xe6, 0x7e, 0x22))

# ===== 五、前端展示设计 =====
add_heading('五、前端展示设计', 1)

add_heading('5.1 新增面板', 2)
add_bullet('患者活动日志面板：实时显示活动切换记录 + 时段摘要按钮')
add_bullet('睡眠质量看板：昨夜睡眠评分 + 近7日趋势图')
add_bullet('医护行为统计面板：巡视频率柱状图 + 访客记录')

add_heading('5.2 与现有系统集成', 2)
add_bullet('活动事件通过MQTT上报，云端新增activity_logs表')
add_bullet('WebSocket实时推送活动切换事件到前端')
add_bullet('交班报告自动包含活动摘要 + 睡眠报告 + 医护记录')

# ===== 六、对赛题评分的加分作用 =====
add_heading('六、对赛题评分的加分作用', 1)
make_table(
    ['评分维度', '如何加分', '分值'],
    [
        ['感知与决策效果', '展示LLM不只处理告警，还能持续感知语义理解', '15分'],
        ['资源与通信效率', '边缘本地生成摘要，只上传文本（vs 上传视频流）', '10分'],
        ['方案完整性', '从"事件驱动"扩展到"持续活动监测"，方案更完整', '15分'],
        ['可扩展性', '活动识别框架可复用到其他场景', '10分'],
        ['应用价值', '护理活动记录是真实临床刚需', '5分'],
        ['创新性', '"边缘LLM持续活动语义汇报"是区别于传统规则告警的创新点', '10分'],
    ],
    [Cm(3.5), Cm(8), Cm(2)]
)

# ===== 七、总结 =====
add_heading('七、总结', 1)
add_para('本次拓展将系统从"被动告警器"升级为"主动感知的AI护理助手"，核心变化：')
add_bullet('从"出事才响应" -> "24小时持续感知"')
add_bullet('从"结构化事件" -> "自然语言护理报告"')
add_bullet('从"单一告警" -> "活动+睡眠+医护多维度"')
add_bullet('从"规则判断" -> "LLM语义理解 + 智能建议"')
add_para('技术实现完全基于现有架构，无需额外硬件，总工时约4天。', bold=True, color=RGBColor(0x1a, 0x56, 0xc4))

out_path = r'C:\Users\Aarontang\Desktop\拓展改进报告-优化排版.docx'
doc.save(out_path)
print('已生成：' + out_path)
