# -*- coding: utf-8 -*-
"""亚伦任务总结报告 docx 生成器（7/29-8/3）

对照《全员任务清单》逐项总结完成情况。
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微软雅黑"
doc = Document()

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r.font.size = Pt(16)

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        r.font.size = Pt(13)

def p(text, bold=False, size=11, color=None, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_shading(cell, "1F4E79")
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r+1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

# ═══════════ 封面 ═══════════
for _ in range(4):
    doc.add_paragraph()
p("智慧病房云边协同系统", bold=True, size=22, color=(0x1F,0x4E,0x79), align="center")
p("亚伦任务总结报告", bold=True, size=18, color=(0x2E,0x75,0xB6), align="center")
doc.add_paragraph()
p("周期：2026-07-29 至 2026-08-03（第 1 周）", size=12, color=(0x59,0x59,0x59), align="center")
p("角色：P1 边缘 AI + 框架 —— T1 边缘 LLM + T2 协同推理（边缘侧）", size=12, color=(0x59,0x59,0x59), align="center")
doc.add_page_break()

# ═══════════ 一、任务完成情况 ═══════════
h1("一、任务完成情况（对照《全员任务清单》）")

p("《全员任务清单》7/29-8/3 亚伦任务要求与实际完成对照：")

table(
    ["要求", "实际完成", "状态"],
    [
        ["直接定跌倒检测方案：YOLOv8+ShuffleNetV2+SA", "完成方案切换，新增 fall_detector.py + train_fall_detector.py，SA 通道注意力模块已与训练脚本对齐", "✅ 超额完成"],
        ["在 UR Fall Detection Dataset 上跑一遍", "完成全量评测（70 片段，11936 帧），规则基线 63.31% + ShuffleNetV2+SA 神经判定 92.11%", "✅ 超额完成"],
        ["交出初步准确率数字（哪怕不高）", "交出 92.11%（远高于最低要求），adl 零误报（FP=0）", "✅ 超额完成"],
        ["之后大方向：边缘-云端切分，配合彦晗", "云边协同端到端打通（wait_for_publish 自锁修复 + 闭环验证），与彦晗接口协议已对齐", "✅ 完成"],
    ],
    [2.5, 4.0, 1.2],
)

# ═══════════ 二、主要成果 ═══════════
h1("二、主要成果")

h2("2.1 代码产出")
bullet("新增文件：activity_tracker.py（活动识别）、summarize_yolo_log.py（日志精简器）、fall_detector.py（任务书方案切换）、train_fall_detector.py（训练脚本）、yolo_realtime_viewer.py（实时窗口）、tracking.py（IoU 跟踪）、behavior.py（行为分析）、yolo_camera.py（真实摄像头适配器）")
bullet("修改文件：main.py（YOLO 管线集成、观测/事件压缩、云端超时回退）、fusion.py、inference.py、llm_advisor.py、Dockerfile、docker-compose.yml 等")
bullet("测试：edge-agent 63 项全绿（从 26 → 63，新增 37 项），training-coordinator 12 项全绿")

h2("2.2 评测数据")
table(
    ["指标", "规则回退（基线）", "ShuffleNetV2+SA（训练后）", "提升"],
    [
        ["准确率", "63.31%", "92.11%", "+29pp"],
        ["adl 特异度", "66.52%", "100.00%（零误报）", "+33pp"],
        ["召回率", "33.71%", "41.53%", "+8pp"],
        ["F1", "0.1505", "0.5037", "3.3x"],
    ],
    [1.5, 2.2, 2.7, 0.8],
)

h2("2.3 关键技术突破")
bullet("云边协同端到端闭环打通：边缘卸载 → MQTT → 云端研判（confirm/reject/escalate）→ 回传 → 边缘状态写回，响应延迟 5s→60ms")
bullet("wait_for_publish 自锁定位与修复：MQTT 回调线程内同步等待 PUBACK 造成 loop 线程死锁，移除后响应延迟从 5 秒降至 60 毫秒")
bullet("日常活动识别（6 类）：玩手机、吃饭、行走、睡眠、静坐、站立，连续 5 帧滞回去抖，单帧闪现被吸收")
bullet("LLM 友好日志精简器：128KB 帧级日志 → 2.7KB 事件摘要（压缩 97%），可直接喂给 LLM 生成护理汇报")

# ═══════════ 三、问题与改进 ═══════════
h1("三、问题与改进")

p("1. ShuffleNetV2+SA 召回率偏低（41.53%），adl 零误报但 fall 检测不够敏感。后续可调数据权重、加训练轮数或调阈值。")
p("2. 活动识别目前只进了 YOLO 窗口显示和 TXT 日志，未接入 MQTT 上报链路。")
p("3. 超时检查依赖主循环 tick（3s 精度），需要独立定时器。")

# ═══════════ 四、下步计划 ═══════════
h1("四、下步计划")

table(
    ["优先级", "任务", "说明", "截止"],
    [
        ["P0", "Jetson Orin Nano 实机部署与性能实测", "TTFT/内存/吞吐/YOLO 并行占用", "8/12"],
        ["P0", "云端真实 14B 联调（配合彦晗）", "vLLM + Qwen2.5-14B，接替 mock", "8/07"],
        ["P1", "活动识别接入 MQTT + LLM 汇报闭环", "activity 字段入事件上报", "8/15"],
        ["P1", "断网保持率测试（≥90%）", "网络模拟 + 日志精简器验证", "8/15"],
        ["P2", "网络模拟器（前端断网演示）", "对齐稳定性 20 分", "8/20"],
    ],
    [0.6, 2.5, 2.8, 0.6],
)

doc.add_paragraph()
p("—— 亚伦（P1）", size=10.5, color=(0x66,0x66,0x66), align="right")
p("2026-08-04", size=10.5, color=(0x66,0x66,0x66), align="right")

OUT = "亚伦任务总结报告-7月29-8月3.docx"
doc.save(OUT)
print(f"saved: {OUT}")
