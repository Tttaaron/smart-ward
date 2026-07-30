# -*- coding: utf-8 -*-
"""优化排版：不足.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== 读取原内容 =====
src = Document(r'C:\Users\Aarontang\Desktop\不足.docx')
original_paras = []
for p in src.paragraphs:
    t = p.text.strip()
    if t:
        original_paras.append({
            'text': p.text,
            'style': p.style.name if p.style else '',
            'bold': any(r.bold for r in p.runs if r.bold),
        })

original_tables = []
for t in src.tables:
    rows = []
    for r in t.rows:
        rows.append([cell.text.strip() for cell in r.cells])
    original_tables.append(rows)

# ===== 辅助函数 =====
def set_cjk(run, name='微软雅黑'):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

# ===== 新建文档 =====
doc = Document()

# 默认样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(3)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('智慧病房云边协同系统\n竞赛差距分析报告')
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
set_cjk(r, '微软雅黑')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(20)
r = sub.add_run('基于挑战杯 XH-202606 赛题要求的逐项评估')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
set_cjk(r, '微软雅黑')

doc.add_page_break()

# ===== 一、已具备的优势 =====
h = doc.add_heading('一、已具备的优势（与赛题契合点）', level=1)
for r in h.runs:
    r.font.size = Pt(16)
    set_cjk(r)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)

# 表格1
t1 = doc.add_table(rows=1, cols=2)
t1.style = 'Table Grid'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t1.rows[0].cells
for cell, text, w, color in zip(hdr, ['赛题要求', '项目现状'], [Cm(4.5), Cm(12)], ['1A56C4', '1A56C4']):
    cell.width = w
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cjk(r)
    set_cell_shading(cell, color)

for req, status in original_tables[0][1:]:
    row = t1.add_row()
    for cell, text, w in zip(row.cells, [req, status], [Cm(4.5), Cm(12)]):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(9.5)
        set_cjk(r)

doc.add_paragraph()

# ===== 二、关键不足 =====
h = doc.add_heading('二、关键不足（按严重程度排序）', level=1)
for r in h.runs:
    r.font.size = Pt(16)
    set_cjk(r)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)

# 跟踪当前严重级别标题
current_section = None
for para in original_paras:
    text = para['text']
    style_name = para['style']
    is_bold = para['bold']

    # 跳过已处理的大标题
    if '一、已具备' in text or '二、关键不足' in text:
        continue

    # 严重级别标题（🔴 🟡 🟢）
    if '致命缺失' in text:
        h = doc.add_heading('🔴 致命缺失（直接不满足硬性指标）', level=2)
        for r in h.runs:
            r.font.size = Pt(13)
            set_cjk(r)
            r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        continue
    if '重要缺失' in text:
        h = doc.add_heading('🟡 重要缺失（影响评分但不致命）', level=2)
        for r in h.runs:
            r.font.size = Pt(13)
            set_cjk(r)
            r.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
        continue
    if '可改进项' in text:
        h = doc.add_heading('🟢 可改进项（加分项缺失）', level=2)
        for r in h.runs:
            r.font.size = Pt(13)
            set_cjk(r)
            r.font.color.rgb = RGBColor(0x27, 0xae, 0x60)
        continue

    # 子标题（如 "1. 无'轻量大模型'"）
    if text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.size = Pt(11.5)
            set_cjk(r)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        continue

    # 普通正文
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.7)
    if '赛题明确要求' in text or '赛题要求的是' in text or '赛题硬指标' in text or '赛题 40 分' in text:
        p.paragraph_format.left_indent = Cm(0.7)
        r = p.add_run(text)
        r.font.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_cjk(r)
        continue
    if '缺失项' in text or '你的项目只有' in text or '你的项目当前' in text or '你的 3 个边缘' in text:
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
        set_cjk(r)
        continue
    if text.startswith('无') or text.startswith('❌') or text.startswith('没有') or text.startswith('缺少'):
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x88, 0x33, 0x33)
        set_cjk(r)
        continue
    if text.startswith('✅'):
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x00, 0x7a, 0x33)
        set_cjk(r)
        continue

    r = p.add_run(text)
    r.font.size = Pt(10)
    set_cjk(r)

doc.add_paragraph()

# ===== 三、量化指标对比表 =====
h = doc.add_heading('三、核心量化指标对比', level=1)
for r in h.runs:
    r.font.size = Pt(16)
    set_cjk(r)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = t2.rows[0].cells
for cell, text, w, color in zip(hdr2, ['指标', '竞赛要求', '项目现状'], [Cm(3), Cm(4), Cm(9.5)], ['1A56C4', '1A56C4', '1A56C4']):
    cell.width = w
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cjk(r)
    set_cell_shading(cell, color)

for metric, req, status in original_tables[1][1:]:
    row = t2.add_row()
    for cell, text, w in zip(row.cells, [metric, req, status], [Cm(3), Cm(4), Cm(9.5)]):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(9.5)
        set_cjk(r)
        if '❌' in text:
            r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

doc.add_paragraph()

# ===== 四、评分预估 =====
h = doc.add_heading('四、评分预估', level=1)
for r in h.runs:
    r.font.size = Pt(16)
    set_cjk(r)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)

scores = [
    ('云边协同技术效果', '40', '10-15', '无大模型压缩、无时延优化、无量化指标'),
    ('方案完整性与可扩展性', '25', '15-18', '架构完整但仅1个场景、训练链路断裂'),
    ('系统稳定性与一致性', '20', '5-8', '无决策冲突检测与仲裁机制'),
    ('创新性与应用价值', '15', '10-12', '病房场景有说服力但方法创新不足'),
]

t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = t3.rows[0].cells
for cell, text, w, c in zip(hdr3, ['评审维度', '满分', '预估得分', '主要卡点'], [Cm(3.5), Cm(1.5), Cm(2), Cm(9.5)], ['1A56C4']*4):
    cell.width = w
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cjk(r)
    set_cell_shading(cell, c)

for dim, full, est, reason in scores:
    row = t3.add_row()
    for cell, text, w in zip(row.cells, [dim, full, est, reason], [Cm(3.5), Cm(1.5), Cm(2), Cm(9.5)]):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(9.5)
        set_cjk(r)
        if dim == '':
            r.font.bold = True

# 总分行
row = t3.add_row()
for cell, text, w in zip(row.cells, ['总分', '100', '40-53', '差距较大，需重点补足大模型压缩、时延优化和决策一致性'], [Cm(3.5), Cm(1.5), Cm(2), Cm(9.5)]):
    cell.width = w
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    set_cjk(r)
    r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

# 保存
out_path = r'C:\Users\Aarontang\Desktop\不足-优化排版.docx'
doc.save(out_path)
print('已优化排版：' + out_path)
