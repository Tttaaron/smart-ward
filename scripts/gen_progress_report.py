# -*- coding: utf-8 -*-
"""项目整体进度报告 docx 生成器"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微软雅黑"
doc = Document()

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r.font.size = Pt(16)

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        r.font.size = Pt(13)

def p(text, bold=False, size=11, color=None, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_shading(cell, "1F4E79")
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r+1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()
p("智慧病房云边协同系统", bold=True, size=22, color=(0x1F,0x4E,0x79), align="center")
p("项目整体进度报告", bold=True, size=18, color=(0x2E,0x75,0xB6), align="center")
doc.add_paragraph()
p("版本：v0.3.0", size=12, color=(0x59,0x59,0x59), align="center")
p("日期：2026 年 7 月 27 日", size=12, color=(0x59,0x59,0x59), align="center")
doc.add_page_break()

# ===== 一、项目概况 =====
h1("一、项目概况")
p("智慧病房云边协同系统采用云-边-端三层架构，复用智能教室项目技术路线（Vue + FastAPI + MQTT + MySQL + SQLite + Docker Compose），领域对象替换为病房安全监护。目标是在每张床位部署独立边缘代理（Orange Pi 5 + RK3588 NPU），本地融合摄像头、床垫、环境三路传感器，实现跌倒/坠床/抽搐等 13 类安全事件的实时识别与闭环处置。")
table(
    ["指标", "数值"],
    [
        ["当前版本", "v0.3.0"],
        ["代码+文档文件", "58 个"],
        ["总行数", "8778 行"],
        ["commit 数", "16 个"],
        ["已合并 PR", "4 个"],
        ["单元测试", "30 项全绿（edge 26 + training 4）"],
    ],
    [2.5, 4.0],
)

# ===== 二、整体完成度 =====
h1("二、整体完成度")
p("项目整体完成度约 55%，各层完成情况如下：")
table(
    ["层面", "完成度", "说明"],
    [
        ["代码层（边缘+云端）", "75%", "边缘端+云端+契约三层扎实，训练空壳"],
        ["文档层", "85%", "技术报告 6/11 章完成，20 份文档"],
        ["硬件层", "15%", "仅选型论证，未到货未实测"],
        ["训练层", "10%", "仅骨架空壳，FedAvg 未实现"],
        ["前端层", "45%", "骨架完成，图表/控制待补"],
    ],
    [2.5, 1.2, 3.3],
)

# ===== 三、各模块完成度 =====
h1("三、各模块完成度")

h2("3.1 边缘端（edge-agent）- 完成度 90%")
table(
    ["功能项", "目标", "实际", "状态"],
    [
        ["采集适配器", "3 类", "Camera/BedSensor/Environment", "✅ 完成"],
        ["融合规则", "11 条", "12 类事件（11 规则 + nurse_call 透传）", "✅ 完成"],
        ["推理引擎接口", "5 字段", "InferenceResult 8 字段 + load_model/rollback", "✅ 完成"],
        ["离线自治", "三件套", "QoS1+SQLite+event_id 去重+补传", "✅ 完成"],
        ["模型下发闭环", "deploy/rollback", "handle_model_deploy + health 上报", "✅ 完成"],
        ["场景驱动", "12 类", "状态机 + 3 床编排", "✅ 完成"],
        ["单元测试", "持续全绿", "26 项", "✅ 完成"],
        ["真实模型接入", "YOLOv8n-pose", "rule-fusion-v1 空壳", "🟡 待硬件"],
    ],
    [1.8, 1.5, 2.8, 1.0],
)

h2("3.2 云端（cloud-backend）- 完成度 90%")
table(
    ["功能项", "实际", "状态"],
    [
        ["数据库", "11 表 ORM + init.sql", "✅"],
        ["REST API", "16 路由（14 业务 + 2 健康）", "✅"],
        ["MQTT 处理器", "4 个 handler（obs/event/health/ack）", "✅"],
        ["WebSocket 推送", "broadcast_sync 桥接 + 心跳", "✅"],
        ["交接班摘要", "generate_shift_summary 自动生成", "✅"],
        ["环境控制", "POST /api/env/control + MQTT 下发", "✅"],
        ["审计日志", "audit_logs 全留痕", "✅"],
        ["node_offline 检测", "无定时扫描任务", "🔴 缺失"],
    ],
    [2.0, 3.5, 1.0],
)

h2("3.3 前端（cloud-frontend）- 完成度 45%")
table(
    ["功能项", "实际", "状态"],
    [
        ["三栏布局", "床位/告警/交接班", "✅"],
        ["WebSocket 实时", "事件推送 + 乐观更新", "✅"],
        ["事件处置 4 动作", "到场/处置/误报/升级", "✅"],
        ["P1 视觉强提示", "红色闪烁动画", "✅"],
        ["床位可视化", "getBedOccupancy 占位函数", "🟡 未接入"],
        ["环境控制 UI", "API 封装但无按钮", "🟡 未接入"],
        ["数据图表", "echarts 依赖已装但无组件", "🔴 缺失"],
        ["组件化", "App.vue 单文件 318 行", "🟡 未拆分"],
    ],
    [2.0, 3.5, 1.0],
)

h2("3.4 协同训练（training-coordinator）- 完成度 25%")
table(
    ["功能项", "实际", "状态"],
    [
        ["调度骨架", "三阶段枚举 + 数据结构", "✅"],
        ["API 空壳", "4 路由", "✅"],
        ["FedAvg 同步（A）", "aggregate 返回占位字符串", "🔴 未实现"],
        ["半异步（B）", "未开始", "🔴"],
        ["稳健剔除（C）", "未开始", "🔴"],
        ["单元测试", "4 项", "✅"],
    ],
    [2.0, 3.5, 1.0],
)

h2("3.5 契约与文档")
table(
    ["项", "实际", "状态"],
    [
        ["MQTT JSON Schema", "6 份 + README", "✅"],
        ["事件枚举", "13 类", "✅"],
        ["严格校验", "additionalProperties: false", "✅"],
        ["基础文档", "13 份（00-12）", "✅"],
        ["技术报告章节", "7 份（骨架+第3/4/5/6/7/10章）", "✅"],
        ["未完成章节", "第1/2/8/9/11章（建鸿/振鑫负责）", "🟡"],
    ],
    [2.0, 3.5, 1.0],
)

# ===== 四、里程碑进度 =====
h1("四、里程碑进度")
table(
    ["里程碑", "截止日期", "目标", "状态"],
    [
        ["M1 需求冻结+架构", "7/15", "方案书+骨架", "✅ 已完成"],
        ["M2 P0 病房闭环", "7/22", "10 功能+联调+测试", "✅ 已完成"],
        ["M3 模型选型+硬件", "7/30", "YOLO 接入+真实数据闭环", "🟡 文档完成，硬件未到"],
        ["M4 P1 云边推理", "8/15", "协同训练+困难样本", "🔴 训练空壳"],
        ["M5 P2 协同训练+测试", "8/25", "半异步+性能看板", "🔴"],
        ["M6 集成+材料+演示", "8/31", "赛事交付", "🔴"],
    ],
    [2.2, 1.0, 2.5, 1.3],
)

# ===== 五、团队各成员进度 =====
h1("五、团队各成员进度")
table(
    ["编号", "姓名", "角色", "主任务", "进度"],
    [
        ["P1", "亚伦", "边缘模型+框架", "模型选型✅/框架✅/报告6章✅/真实模型🟡", "90%"],
        ["P2", "景彬", "场景+前端", "scenario✅/前端骨架🟡", "50%"],
        ["P3", "建鸿", "统筹+训练+方案", "赛事对接🟡/FedAvg🔴/方案定稿🔴", "20%"],
        ["P4", "振鑫", "协同训练底层", "骨架✅/FedAvg🔴/半异步🔴", "25%"],
        ["P5", "先伟", "扩散模型调优+数据集", "未开始🔴", "0%"],
        ["P6", "烽亮", "扩散模型开发+视频", "未开始🔴", "0%"],
        ["P7", "彦晗", "云边模型协同", "未开始🔴", "0%"],
    ],
    [0.5, 0.8, 1.5, 2.7, 0.7],
)

# ===== 六、关键缺口与风险 =====
h1("六、关键缺口与风险")

h2("6.1 P0 关键路径阻塞")
table(
    ["缺口", "影响", "负责人"],
    [
        ["FedAvg 同步基线未实现", "云边协同创新点无法演示", "振鑫"],
        ["扩散模型未开始", "困难样本生成闭环缺失", "烽亮/先伟"],
        ["真实模型接入 inference.py", "边缘识别仍是模拟", "亚伦（待硬件）"],
        ["node_offline 云端定时检测", "断网演示场景不完整", "建鸿/景彬"],
    ],
    [3.0, 3.0, 1.0],
)

h2("6.2 P1 功能缺口")
table(
    ["缺口", "影响", "负责人"],
    [
        ["前端床位可视化占位", "getBedOccupancy 未接入", "景彬"],
        ["前端环境控制无按钮", "triggerEnvControl 未接入UI", "景彬"],
        ["前端无 echarts 图表", "趋势图/延迟看板缺失", "景彬"],
        ["前端未组件化", "App.vue 单文件", "景彬"],
    ],
    [3.0, 3.0, 1.0],
)

h2("6.3 风险评估")
bullet("训练算法（FedAvg）和扩散模型两部分基本是空的，这是项目云边协同创新点的核心，但 P3-P7 五人进度明显滞后于亚伦的 P1。")
bullet("距离 8/31 赛事交付约 45% 工作量待完成，关键路径是训练算法（振鑫）+ 扩散模型（烽亮/先伟）+ 真实模型接入（亚伦待硬件）。")
bullet("硬件到货延迟风险：需求 §10.1 风险表已识别，应对方案是先用模拟数据，到货后 5 天落地。")

doc.save("项目整体进度报告.docx")
print("✓ 项目整体进度报告.docx 生成成功")
