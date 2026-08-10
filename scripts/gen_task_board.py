# -*- coding: utf-8 -*-
"""全员任务清单看板生成器（2026-08-05 版）。

依据《项目.docx》执行版任务书 + 8/5 最新现状生成，按成员分区，
每人的"已完成内容"列可直接填写（命名规范：日期_模块_场景_提交号）。
输出到桌面：全员任务清单看板-20260805.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# 默认样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(3)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '微软雅黑')


def set_cjk(run, name='微软雅黑'):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_heading(text, level=1, color='1A56C4'):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_cjk(r)
        r.font.color.rgb = RGBColor(*[int(color[i:i + 2], 16) for i in (0, 2, 4)])
    return h


def add_para(text, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    set_cjk(r)
    return p


def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def status_color(status):
    if status.startswith('✅'):
        return 'C6EFCE'   # 绿
    if status.startswith('🟡'):
        return 'FFEB9C'   # 黄
    if status.startswith('❌'):
        return 'FFC7CE'   # 红
    if status.startswith('⬜'):
        return 'D9D9D9'   # 灰
    return None


def member_table(member_rows):
    """每人一张表：任务项 | 截止 | 状态 | 已完成内容（待填写）"""
    widths = [Cm(7.2), Cm(1.6), Cm(1.6), Cm(5.8)]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text, w in zip(table.rows[0].cells, ['任务项', '截止', '状态', '已完成内容（请填写）'], widths):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cjk(r)
        set_cell_shading(cell, '1A56C4')
    for task, deadline, status, done in member_rows:
        row = table.add_row()
        for cell, text, w in zip(row.cells, [task, deadline, status, done], widths):
            cell.width = w
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(9)
            set_cjk(r)
            if status == '✅ 完成':
                r = cell.paragraphs[0].add_run('')
        color = status_color(status)
        if color:
            set_cell_shading(row.cells[2], color)
    doc.add_paragraph()


def snapshot_table(rows):
    """总体快照表（指标 | 内容）"""
    widths = [Cm(3.6), Cm(12.6)]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text, w in zip(table.rows[0].cells, ['指标', '现状（2026-08-05）'], widths):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cjk(r)
        set_cell_shading(cell, '1A56C4')
    for metric, content in rows:
        row = table.add_row()
        for cell, text, w in zip(row.cells, [metric, content], widths):
            cell.width = w
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(9)
            set_cjk(r)
    doc.add_paragraph()


def member_section(no, name, role, status_line, rows):
    add_heading(f'{no} {name}（{role}）', 2, '2E75B6')
    add_para('当前状态：' + status_line, size=9.5, color=RGBColor(0x59, 0x59, 0x59))
    member_table(rows)


# ===== 封面 =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('全员任务清单看板')
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
set_cjk(r)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('智慧病房云边端协同系统 · 2026-08-05 版 · 最终截止 8/31')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
set_cjk(r)

doc.add_paragraph()

# ===== 使用说明 =====
add_heading('使用说明', 1)
for line in [
    '1. 看自己名字那一段就够了：每人的"已完成内容"列直接填写，写清做了什么、证据在哪（命令/日志/截图/提交号）。',
    '2. 状态标记：✅ 完成 / 🟡 进行中 / ❌ 未开始 / ⬜ 阻塞。每周联调前更新一次，8/31 冻结。',
    '3. 证据命名规范：日期_模块_场景_提交号，如 20260812_edge-llm_jetson-1.5b_<commit>。所有数字必须回指原始证据。',
    '4. 已完成内容由 P3 建鸿按验收矩阵抽查，未达标项不得标为"完成"。',
]:
    add_para(line, size=9.5)

# ===== 总体快照 =====
add_heading('总体进度快照（2026-08-05）', 1)
snapshot = [
    ['测试基线', 'edge-agent 78 项、training-coordinator 12 项、cloud-llm-service 9 项，全绿'],
    ['跌倒检测', 'UR 全量评测：帧级准确率 95.78%、召回 77.50%、F1 0.78；片段级 100%（评测口径 bug 已修复）'],
    ['云边协同', 'cloud-llm-service mock 消费者就绪（归先伟）；独立超时定时器、契约 Schema、活动识别上报已合入 master'],
    ['训练线', 'FedBuff 异步聚合 + MiniLLM/Hinton 蒸馏已提交（dddfcf3）'],
    ['最大缺口', 'Jetson 实机验收（等 P6 设备确认，8/12 截止）；云端真实 14B 接入（8/15 前）；断网保持率测试（8/15）'],
]
snapshot_table(snapshot)

# ===== 成员分区 =====
add_heading('个人任务清单（按名字区分）', 1)

member_section('P1', '亚伦', '边缘 AI + 云边协同（边缘侧）',
    '边缘主线基本完成：跌倒评测修复（95.78%）、超时定时器、活动识别上报、bench 脚本、Schema 均已提交。剩 Jetson 实测与联调取证。',
    [
        ('Jetson 双路径实机验收（1.5B/0.5B：TTFT<200ms、RSS≤1.5GB、吞吐、YOLO 并行占用）', '8/12', '🟡 进行中',
         'bench_jetson.py 已就绪（已提交）；设备确认待 P6 答复'),
        ('云端超时独立定时机制（解耦 TICK_SECONDS）', '8/7', '✅ 完成',
         'main.py 守护线程 + 3 测试（7348d7a）'),
        ('inference request/response Schema + 字段校验', '8/8', '✅ 完成',
         'contracts/inference_request.json + response.json + 7 契约测试（5b22b35），待与先伟会签'),
        ('与先伟真实 Broker 联调 7 场景（正常/超时/重复/非法/未知/trace/断网，每场景≥20次）', '8/11', '🟡 进行中',
         'cloud-llm-service mock 已就绪，可随时联调'),
        ('跌倒检测评测（UR Fall Detection Dataset）', '8/12', '✅ 完成',
         '评测口径修复：准确率 95.78%、召回 77.50%（9278eb0）；证据 docs/evidence/'),
        ('活动识别接入 MQTT 上报链路', '8/15', '✅ 完成',
         'observation.activity + 事件 details 透传（41a0f2f）'),
        ('活动 LLM 汇报：模式A 实时播报 + 模式B 时段摘要', '8/15', '❌ 未开始',
         '——待填写'),
        ('状态一致性：可追踪 trace_id 的完整端到端日志', '8/15', '🟡 进行中',
         '——待填写'),
    ])

member_section('P2', '景彬', '护士站前端 + 可观测性 + 测试数据',
    '前端可观测性看板已完成（路由/性能/网络/状态 + 心跳 + 截图标注 + 录屏）。剩活动面板、异常态素材与 MIMIC-IV 认证跟进。',
    [
        ('npm run build + Docker 镜像构建验证', '8/7', '✅ 完成',
         '构建日志 docs/evidence/20260804_frontend_*.log（f58d27d）'),
        ('路由/性能/网络/状态可观测性看板', '8/10', '✅ 完成',
         'SystemStatusBar + EventDetailDrawer + eventMeta（f58d27d/bc5ed11）'),
        ('活动日志面板（对接 observation.activity 字段）', '8/15', '🟡 进行中',
         '边缘 activity 字段已就绪（41a0f2f），面板待开发——待填写'),
        ('WebSocket 异常态演示素材（断网/重连/补传，trace_id 标注）', '8/12', '🟡 进行中',
         '已有 MQTT 场景截图与 webm 录屏（bc5ed11）——待填写'),
        ('MIMIC-IV/eICU 资格认证申请与体征模拟器', '8/20', '🟡 进行中',
         '认证申请已提交（7/29 周）——状态待填写'),
        ('模拟传感器代码替换（模拟真实硬件）', '8/20', '🟡 进行中',
         '——待填写'),
    ])

member_section('P3', '建鸿', '统筹 + 训练核心开发 + 方案材料',
    'FedBuff 异步聚合 + MiniLLM/Hinton 蒸馏已提交（dddfcf3）；统筹线待落地：验收矩阵、材料一致性、周联调。',
    [
        ('分布式协同训练：梯度聚合 + 同步异步调度代码（含振鑫/彦晗）', '8/22', '✅ 完成',
         'FedBuff async aggregation + MiniLLM/Hinton distillation（dddfcf3）——证据待填写'),
        ('发布验收矩阵（指标/目标/实现位置/命令/证据/负责人/状态/风险）', '8/3', '🟡 进行中',
         '——待填写'),
        ('统一材料口径：README/CHANGELOG/报告数字一致，8/4 基线快照', '8/15', '🟡 进行中',
         '——待填写'),
        ('每周联调会议（8/7、8/15、8/22、8/28）纪要 + 问题闭环', '每周', '🟡 进行中',
         '——待填写'),
        ('赛事对接：官方群/材料明细/权重结论归档', '8/10', '🟡 进行中',
         '——待填写'),
        ('最终材料：技术报告/PPT/提交包冻结', '8/28', '❌ 未开始',
         '——待填写'),
    ])

member_section('P4', '振鑫', '协同训练底层调度 + 一致性',
    'FedBuff 异步聚合与蒸馏已落地（配合建鸿）；版本一致性、回滚演示证据待补。',
    [
        ('FedBuff 异步聚合 + 陈旧度加权实现', '8/22', '✅ 完成',
         'dddfcf3（与建鸿/彦晗协作）——证据待填写'),
        ('MiniLLM/Hinton 蒸馏（14B 教师 → 1.5B/0.5B 学生）', '8/22', '✅ 完成',
         'dddfcf3——实验日志待填写'),
        ('模型版本关系表（节点/聚合/教师/学生/发布批次命名规则）', '8/15', '🟡 进行中',
         '——待填写'),
        ('版本冲突/迟到节点/重复上传处理规则 + 测试', '8/15', '🟡 进行中',
         '——待填写'),
        ('灰度发布/失败回滚演示 + 边缘 health 确认记录', '8/20', '❌ 未开始',
         '——待填写'),
    ])

member_section('P5', '先伟', '云端 LLM 服务 + 扩散模型调优 + 测试统筹',
    'cloud-llm-service mock 消费者完成（9 测试全绿），归属先伟。扩散调优与 NLU 数据集进行中。',
    [
        ('cloud-llm-service：mock 消费者 + 健康检查 + 去重/幂等', '8/7', '✅ 完成',
         'main.py/mqtt_handler/llm_client + 9 测试（74a5082/2f0d991）——证据待填写'),
        ('云端真实 Qwen2.5-14B/vLLM 接入（替换 mock）', '8/15', '🟡 进行中',
         '——待填写'),
        ('与亚伦真实 Broker 联调 7 场景 + Schema 会签', '8/11', '🟡 进行中',
         '——待填写'),
        ('云端扩散模型调优 + 困难样本筛选（SHAP 引导 + 条件扩散）', '8/25', '🟡 进行中',
         '——待填写'),
        ('NLU 评测集（500+ 条，含云端处置判断标签）', '8/25', '🟡 进行中',
         'docs/20-NLU数据集构建说明-P5.md 已建——数据量待填写'),
        ('断网保持率测试（≥90%，三指标分开统计）', '8/15', '❌ 未开始',
         '——待填写'),
        ('三路线对比报告（1.5B/0.5B/14B：准确率/F1/TTFT/RSS/吞吐）', '8/25', '❌ 未开始',
         '——待填写'),
    ])

member_section('P6', '烽亮', '扩散模型开发部署 + 演示视频',
    'diffusion-service 已有完整代码（generator/curator/exporter + API）。Compose 构建与视频待验收；Jetson 环境答复待确认。',
    [
        ('云端扩散模型整体开发 + 部署落地（困难样本生成/数据集扩充）', '8/25', '🟡 进行中',
         'diffusion-service 服务代码已存在（fa38083）——部署日志待填写'),
        ('Jetson Orin Nano 环境准备（CUDA/模型挂载/日志目录/监控命令）', '8/12', '⬜ 阻塞',
         '待确认设备可用日期（8/5 答复亚伦）'),
        ('主 Compose / compact Compose / 前端镜像构建验证', '8/15', '🟡 进行中',
         '——待填写'),
        ('云端 14B/vLLM 运行环境（先伟接入的依赖）', '8/15', '🟡 进行中',
         '——待填写'),
        ('5~8 分钟演示脚本 + 最终视频（正常/云端/断网/恢复/路径切换）', '8/28', '🟡 进行中',
         '已有前端演示 webm 素材（bc5ed11）——完整视频待填写'),
    ])

member_section('P7', '彦晗', '协同训练技术开发（配合建鸿）',
    '已参与 FedBuff 异步聚合与蒸馏提交。与振鑫的分工边界待细化。',
    [
        ('分布式协同训练：梯度聚合 + 同步异步调度（配合建鸿/振鑫）', '8/22', '✅ 完成',
         'dddfcf3——个人分工部分待填写'),
        ('与振鑫分工确认（调度算法/一致性模块划分）', '8/10', '🟡 进行中',
         '——待填写'),
        ('训练调度实验日志 + 数据归档', '8/25', '🟡 进行中',
         '——待填写'),
    ])

# ===== 共享里程碑 =====
add_heading('共享里程碑（全员）', 1)
member_table([
    ['云端最小闭环：一条 request 经 Broker 到 response（P5 主责）', '8/7', '🟡 进行中', 'mock 已就绪——联调记录待填写'],
    ['云边真实联调：7 场景取证（亚伦 + 先伟）', '8/11', '🟡 进行中', '——待填写'],
    ['Jetson 双路径验收（亚伦 + 烽亮 + 先伟）', '8/12', '⬜ 阻塞', '等设备确认'],
    ['断网保持率 ≥90% 测试（先伟 + 亚伦）', '8/15', '❌ 未开始', '——待填写'],
    ['前端演示闭环（景彬 + 亚伦）', '8/20', '🟡 进行中', '——待填写'],
    ['精度/稳定性/训练实验证据（振鑫 + 先伟）', '8/25', '🟡 进行中', '——待填写'],
    ['材料初稿 + 全链路彩排（P3 统筹）', '8/28', '❌ 未开始', '——待填写'],
    ['最终提交（8/31 23:59 前）', '8/31', '❌ 未开始', '——待填写'],
])

# ===== 保存 =====
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
out_path = os.path.join(desktop, '全员任务清单看板-20260805.docx')
doc.save(out_path)
print('已生成：' + out_path)
