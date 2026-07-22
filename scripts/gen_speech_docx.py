# -*- coding: utf-8 -*-
"""演讲稿 docx 生成器（简洁版，5-8 分钟）"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微软雅黑"
doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def set_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r.font.size = Pt(16)

def p(text, bold=False, size=12, color=None, align=None, indent=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.6
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def cue(text):
    """演讲提示（灰色斜体）"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run("【提示】" + text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_shading(cell, "1F4E79")
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r+1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

# ===== 封面 =====
for _ in range(5):
    doc.add_paragraph()
p("智慧病房云边协同系统", bold=True, size=24, color=(0x1F,0x4E,0x79), align="center")
p("演讲稿", bold=True, size=20, color=(0x2E,0x75,0xB6), align="center")
doc.add_paragraph()
p("汇报人：亚伦", size=13, color=(0x40,0x40,0x40), align="center")
p("2026 年 7 月", size=12, color=(0x59,0x59,0x59), align="center")
doc.add_page_break()

# ===== 开场 =====
h1("一、开场")
p("各位老师、各位同学，大家好。我是亚伦，今天向大家汇报智慧病房云边协同系统的进展。", size=13)
p("这个项目脱胎于我们之前已经验证过的智能教室云边协同节能系统。我们复用了它已经成熟的技术路线——Vue、FastAPI、MQTT、MySQL、Docker Compose——但把领域模型、数据契约、事件闭环、边缘推理，全部面向智慧病房场景重新设计。", size=13)
p("一句话定位：这是一个患者安全辅助系统，用于赛事展示和原型演示，不替代医护人员处置。所有高风险事件都由护士确认后才处置。", size=13)
cue("语速放慢，强调“患者安全辅助”定位，划清与医疗器械的边界。")

# ===== 二、为什么做这个项目 =====
h1("二、为什么做这个项目")
p("病房里有一个很现实的矛盾：护士看不过来，跑不过来，记不过来。一个护士要管多张床位，夜间巡房靠走动，交接班靠手写，事件发生后才知道。", size=13)
p("我们想解决三个问题：", size=13)
p("第一，把“事后发现”变成“事前预警”。比如坠床，传统是摔了才知道；我们希望在患者翻身到床沿时就预警。", size=13, indent=0.3)
p("第二，把“护士巡房”变成“事件驱动”。只有异常才推送告警，护士不用反复走动。", size=13, indent=0.3)
p("第三，把“人盯人”变成“系统留痕”。每个事件都有时间、置信度、确认人、处置结果，事后可追溯。", size=13, indent=0.3)
cue("用病房真实痛点切入，不要先讲技术。强调“辅助”而非“替代”。")

# ===== 三、系统架构 =====
h1("三、系统架构")
p("系统采用云边端三层协同架构：", size=13)
table(
    ["层级", "职责", "特点"],
    [
        ["边缘端", "采集 + 轻量推理 + 即时告警 + 本地缓存", "1 秒内告警，断网可用"],
        ["云端", "事件中心 + 模型训练 + 版本下发 + 数据分析", "复杂复核，不阻塞紧急告警"],
        ["护士站", "实时告警 + 确认处置 + 交接班摘要", "WebSocket 增量推送"],
    ],
    [1.0, 3.0, 2.0],
)
p("核心设计原则：紧急事件边缘端本地决策，不等待云端；灰区事件上传脱敏片段供云端复核；断网时继续采集告警，恢复后补传不重复。", size=13)
cue("这里可以指着架构图讲，强调“边缘不等待云端”是核心创新点。")

# ===== 四、已交付成果 =====
h1("四、已交付成果")
p("目前项目整体技术框架已完成端到端联调验证，具体成果：", size=13)

h1("4.1 代码规模")
table(
    ["模块", "内容", "规模"],
    [
        ["contracts/", "6 份 MQTT JSON Schema（联调标准）", "6 文件"],
        ["edge-agent/", "4 类采集适配器 + 融合引擎 + 推理接口 + 场景驱动", "8 文件"],
        ["cloud-backend/", "11 张表 + 21 个 REST API + WebSocket", "7 文件"],
        ["cloud-frontend/", "Vue 3 护士站工作台", "4 文件"],
        ["training-coordinator/", "协同训练调度空壳", "2 文件"],
    ],
    [1.8, 4.0, 1.2],
)
p("代码约 4970 行，8 个 commit，17 项单元测试全绿。", bold=True, size=13)

h1("4.2 10 项智能功能")
p("在框架基础上实现了 10 项病房智能功能，覆盖三类场景：", size=13)
table(
    ["类别", "功能", "亮点"],
    [
        ["患者安全", "坠床预警（事前）、跌倒检测、抽搐检测、长时间静止、异常体态、压疮预防", "坠床预警是事前，不是事后"],
        ["护理响应", "交接班摘要自动生成", "按班次聚合 12 事件含分布"],
        ["环境设备", "环境自适应、空气质量联动、床位占用可视化、设备故障预警", "CO₂ 超阈值自动开新风"],
    ],
    [1.2, 3.5, 2.3],
)
cue("重点讲“坠床预警”和“交接班摘要”，这两个是差异化亮点。")

h1("4.3 端到端联调验证")
p("8 个 Docker 服务全部启动并通过联调：", size=13)
p("3 个边缘代理实时上报观测与健康心跳，3 床在线；", size=13, indent=0.3)
p("6 类新事件自动触发并入库（坠床预警/长时间静止/异常体态/抽搐/压疮/设备故障）；", size=13, indent=0.3)
p("告警闭环验证：事件 new -> notified -> acknowledged -> resolved，含处置记录与审计日志；", size=13, indent=0.3)
p("交接班摘要 API 生成 12 事件统计含类型分布；", size=13, indent=0.3)
p("环境控制指令通过 MQTT 下发到边缘端。", size=13, indent=0.3)
p("联调中发现并修复 4 个真实 bug，均已提交。", size=11, color=(0x59,0x59,0x59))

# ===== 五、关键技术点 =====
h1("五、关键技术点")

h1("5.1 复用而非推倒重来")
p("我们从智能教室项目复用了 8 个文件（WebSocket 管理器、日志、Dockerfile、前端代理等），模式复用 + 字段重写 4 个文件（数据库、MQTT 处理器、主程序、Schema），只新写了病房特有的部分（11 表 SQL、新 MQTT 主题树、4 类适配器、融合引擎）。", size=13)
p("这样既保留了已验证的稳定性，又快速适配了病房场景。", size=13)

h1("5.2 事件状态机")
p("我们把“识别事件”“通知任务”“人工处置”分开建模。事件生命周期：new -> notified -> acknowledged -> resolved / false_positive / escalated。", size=13)
p("每个事件都带置信度、模型版本、证据引用，护士标记的误报可以回流作为训练负样本。", size=13)

h1("5.3 断网自治")
p("这是云边架构相比纯云端的最大优势。网络中断时：边缘端继续采集推理告警；事件写入本地 SQLite 标记未同步；网络恢复后按序号补传，云端按 event_id 去重，不丢不重。", size=13)
cue("这里可以现场演示断网：停掉某个 edge-agent 容器的网络，事件仍本地缓存，恢复后补传。")

h1("5.4 隐私保护")
p("原始视频留边缘端，默认不长期保存；仅上传事件摘要、脱敏截图；前端用匿名别名（张阿姨/李伯伯），不存真实姓名；人脸模糊默认开启。", size=13)

# ===== 六、演示 =====
h1("六、现场演示")
p("现在打开浏览器访问 http://localhost:8081，大家可以看到护士站工作台：", size=13)
p("顶栏显示病区 1、床位 3、节点 3/3 在线、P1 待处理告警数；", size=13, indent=0.3)
p("左侧床位卡片显示患者别名与状态；", size=13, indent=0.3)
p("中间告警工作台实时推送事件，可点击到场/处置/误报/升级；", size=13, indent=0.3)
p("右侧交接班摘要面板，选班次点生成，自动汇总本班事件。", size=13, indent=0.3)
cue("演示节奏：先看顶栏统计 -> 看告警弹出 -> 点击到场 -> 生成交接班摘要。约 2 分钟。")

# ===== 七、硬件与部署 =====
h1("七、硬件与部署方案")
p("当前框架已验证模拟数据闭环。下一步接入真实硬件，1 床位最小演示约 1400 元：", size=13)
table(
    ["硬件", "型号", "价格", "用途"],
    [
        ["边缘计算盒", "Orange Pi 5（8GB，RK3588 NPU）", "800 元", "运行 edge-agent + 模型推理"],
        ["USB 摄像头", "罗技 C920", "400 元", "姿态识别"],
        ["床垫传感器", "FSR402 + ESP32", "100 元", "占床/离床检测"],
    ],
    [1.3, 2.5, 0.8, 2.4],
)
p("接入真实硬件时不改现有代码框架，只需替换 3 个适配器的 read() 方法。5 天可落地 1 床位真实数据闭环。", size=13)

# ===== 八、未完成与下一步 =====
h1("八、未完成与下一步")
p("当前任务1（边缘模型选型）完成 30%，任务2（场景调研）完成 40%。下一步：", size=13)
p("近期：产出 YOLOv8n-pose / OpenVINO / TensorRT 模型选型对比文档；", size=13, indent=0.3)
p("近期：在 inference.py 接入真实 YOLO 模型，目标硬件实测延迟与精度；", size=13, indent=0.3)
p("近期：采购硬件，5 天落地 1 床位真实演示；", size=13, indent=0.3)
p("中期：配合建鸿完成完整项目调研，配合团队接入协同训练与扩散模型；", size=13, indent=0.3)
p("8/31 前：完成赛事申报材料技术部分。", size=13, indent=0.3)

# ===== 九、总结 =====
h1("九、总结")
p("最后总结三句话：", size=13, bold=True)
p("第一，技术框架已交付并端到端验证通过，8 服务跑通，10 项智能功能落地，17 项测试全绿。", size=13, indent=0.3)
p("第二，架构核心是云边协同——边缘端自治保证实时性，云端复核与训练保证持续进化。", size=13, indent=0.3)
p("第三，下一步聚焦真实硬件接入与模型选型，5 天内跑通 1 床位真实数据闭环。", size=13, indent=0.3)
p("汇报完毕，谢谢大家。", size=13, bold=True)
cue("鞠躬，停顿，等待提问。预计提问方向：模型选型、隐私合规、真实部署成本、断网演示。")

# ===== 保存 =====
out = "docs/演讲稿.docx"
doc.save(out)
import os
print(f"✅ 已生成：{out} ({os.path.getsize(out)/1024:.1f} KB)")
