# -*- coding: utf-8 -*-
"""生成 training-coordinator 代码审查报告 docx"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cjk_font(run, font_name='微软雅黑'):
    """设置 run 的中文字体（确保 rPr 存在）"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(4)

# 封面标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(80)
run = title.add_run('协同训练调度层代码审查报告')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
set_cjk_font(run)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('training-coordinator 模块 Bug 分析')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
set_cjk_font(run)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.paragraph_format.space_before = Pt(40)
today = datetime.date.today().strftime('%Y-%m-%d')
run = info.add_run('审查日期：' + today + '\n审查范围：training-coordinator/\n测试结果：8 项单元测试全部通过')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
set_cjk_font(run)

doc.add_page_break()


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_cjk_font(r)
        if level == 1:
            r.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
        elif level == 2:
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h


def add_para(text, bold=False, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    set_cjk_font(r)
    return p


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    return p


def sev_color(s):
    m = {'阻断': RGBColor(0xc0, 0x39, 0x2b),
         '严重': RGBColor(0xe6, 0x7e, 0x22),
         '一般': RGBColor(0x27, 0xae, 0x60)}
    return m.get(s, RGBColor(0x33, 0x33, 0x33))


# 一、概述
add_heading('一、审查概述', 1)
add_para('本次审查针对 training-coordinator 模块的全部代码（app/scheduler.py、app/main.py、demo/、tests/），结合单元测试运行和静态代码分析，共发现 10 个问题，其中阻断级 3 个、严重级 3 个、一般级 4 个。', indent=True)
add_para('值得注意的是，8 项单元测试全部通过（pytest 0.23s），但测试仅覆盖了算法层（FedAvgScheduler / SemiAsyncScheduler）的独立逻辑，未覆盖编排层与算法层的集成路径，也未覆盖 numpy 2.x 运行环境，因此 3 个阻断级 bug 未被测试发现。', indent=True)

# 二、问题汇总表
add_heading('二、问题汇总表', 1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
headers = ['#', '问题描述', '严重程度', '文件位置', '负责人']
widths = [Cm(0.8), Cm(7), Cm(1.8), Cm(4), Cm(1.8)]
for cell, text, w in zip(hdr, headers, widths):
    cell.width = w
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cjk_font(r)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1A56C4')
    tcPr.append(shd)

bugs = [
    ('1', '编排层用假权重占位，真实梯度从未接入', '阻断', 'app/scheduler.py:113-116', '振鑫'),
    ('2', 'Demo 依赖的 src/ 目录缺失，无法运行', '阻断', 'demo/run_fedavg.py:15-21', '振鑫'),
    ('3', 'np.zeros_like numpy 2.x 不兼容', '阻断', 'app/scheduler.py:216', '振鑫'),
    ('4', 'ClientUpdate 类型声明与实际使用不一致', '严重', 'app/scheduler.py:46', '振鑫'),
    ('5', 'run_semi_async.py 缺少 import numpy', '严重', 'demo/run_semi_async.py:37', '振鑫'),
    ('6', 'SemiAsync 陈旧度始终为0，加权失效', '严重', 'app/scheduler.py:285-286', '振鑫'),
    ('7', 'start_round 参数解析不安全', '一般', 'app/main.py:30', '建鸿'),
    ('8', 'requirements.txt 缺少 numpy 依赖', '一般', 'requirements.txt', '建鸿'),
    ('9', 'SemiAsync 未按样本数加权', '一般', 'app/scheduler.py:296-327', '振鑫'),
    ('10', '_algo_engine 缓存导致跨轮污染', '一般', 'app/scheduler.py:135-144', '振鑫'),
]
for row_data in bugs:
    row = table.add_row()
    for i, (cell, text, w) in enumerate(zip(row.cells, row_data, widths)):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(9.5)
        if i == 2:
            r.font.bold = True
            r.font.color.rgb = sev_color(text)
        set_cjk_font(r)

doc.add_paragraph()

# 三、阻断级
add_heading('三、阻断级问题详情', 1)

add_heading('Bug 1：编排层用假权重占位，真实梯度从未接入', 2)
add_para('文件：app/scheduler.py 第 113-116 行', bold=True)
add_para('问题描述：', bold=True)
add_para('TrainingScheduler.aggregate() 方法在向算法层传递客户端更新时，没有使用边缘端实际上报的梯度数据，而是用 _dummy_weight_template() 生成的硬编码空白张量替代。所有客户端的权重统一被替换为 [zeros(2,3), zeros(2,)]，导致算法层 FedAvgScheduler.aggregate() 聚合的是一堆零矩阵，聚合结果毫无意义。', indent=True)
add_para('相关代码：', bold=True)
add_code('dummy_weights = _dummy_weight_template()\nclient_updates[i] = (dummy_weights, upd.sample_count)')
add_para('影响：编排层与算法层之间的权重传递完全断裂，协同训练链路跑不通。', bold=True, color=RGBColor(0xc0, 0x39, 0x2b))

add_heading('Bug 2：Demo 依赖的 src/ 目录缺失', 2)
add_para('文件：demo/run_fedavg.py 第 15-21 行', bold=True)
add_para('问题描述：两个 demo 脚本都通过 from src.scheduler / from src.models / from src.utils 导入模块，但 src/ 目录在仓库中不存在。这些文件是振鑫本地的测试工具库，提交时未带上。', indent=True)
add_para('影响：两个 demo 均无法运行，无法验证算法收敛性。', bold=True, color=RGBColor(0xc0, 0x39, 0x2b))

add_heading('Bug 3：np.zeros_like numpy 2.x 不兼容', 2)
add_para('文件：app/scheduler.py 第 216 行', bold=True)
add_para('问题描述：FedAvgScheduler.aggregate() 中使用 np.zeros_like(w) 遍历异构形状的 ndarray 列表。在 numpy 2.x（当前环境 numpy 2.5）下会抛出 ValueError: inhomogeneous shape。实测在当前环境直接崩溃。', indent=True)
add_para('相关代码：', bold=True)
add_code('aggregated = [np.zeros_like(w) for w in first_w]')
add_para('修复建议：改用 np.zeros(w.shape, dtype=w.dtype)', indent=True)

# 四、严重级
add_heading('四、严重级问题详情', 1)

add_heading('Bug 4：ClientUpdate 类型声明与实际使用不一致', 2)
add_para('文件：app/scheduler.py 第 46 行', bold=True)
add_para('问题描述：ClientUpdate.weights_summary 声明为 Dict[str, float]，但算法层实际操作的是 List[np.ndarray]。main.py 提交更新时传的是空字典 weights_summary={}，真实权重无处可放。', indent=True)

add_heading('Bug 5：run_semi_async.py 缺少 import numpy', 2)
add_para('文件：demo/run_semi_async.py 第 37 行', bold=True)
add_para('问题描述：main() 函数内第 37 行使用了 np.concatenate()，但 import numpy as np 被写在 if __name__ 块内（第 109 行），函数执行时 np 未定义，会抛出 NameError。', indent=True)

add_heading('Bug 6：SemiAsync 陈旧度始终为0，加权形同虚设', 2)
add_para('文件：app/scheduler.py 第 285-286 行', bold=True)
add_para('问题描述：receive_update() 中 staleness = self.round - client_round。但 simulate_concurrent_round() 中所有客户端都用同一个 current_round，且 self.round 在聚合后才 +1。导致同一轮内提交的所有客户端 staleness 始终为 0，陈旧度加权公式 1/(staleness+1) 恒等于 1.0，加权机制完全失效。', indent=True)

# 五、一般级
add_heading('五、一般级问题详情', 1)

add_heading('Bug 7：start_round 参数解析不安全', 2)
add_para('文件：app/main.py 第 30 行', bold=True)
add_para('问题描述：participants: list[str] 作为 FastAPI 参数会被尝试从 query string 解析，前端传 JSON body 会失败。应使用 Pydantic 模型定义请求体。', indent=True)

add_heading('Bug 8：requirements.txt 缺少 numpy 依赖', 2)
add_para('文件：training-coordinator/requirements.txt', bold=True)
add_para('问题描述：scheduler.py 第 25 行 import numpy as np，但 requirements.txt 未声明 numpy。Docker 构建时依赖系统全局 numpy，声明不完整。', indent=True)

add_heading('Bug 9：SemiAsync 未按样本数加权', 2)
add_para('文件：app/scheduler.py 第 296-327 行', bold=True)
add_para('问题描述：_aggregate_pending() 聚合时只用陈旧度权重 1/(staleness+1)，没有乘以样本数 num_samples。而 FedAvgScheduler 是按样本数加权的。两个策略的聚合口径不一致。', indent=True)

add_heading('Bug 10：_algo_engine 缓存导致跨轮污染', 2)
add_para('文件：app/scheduler.py 第 135-144 行', bold=True)
add_para('问题描述：_get_algo_engine() 首次创建 _algo_engine 后永久复用。但不同轮次的 n_clients 和 init_weights 可能不同，缓存的引擎沿用旧的客户端数和初始权重，导致后续轮次聚合异常。', indent=True)

# 六、测试覆盖分析
add_heading('六、测试覆盖分析', 1)
add_para('当前 8 项单元测试全部通过，但覆盖范围存在明显盲区：', indent=True)
gaps = [
    '编排层与算法层的集成路径未测试（Bug 1 未被发现）',
    'Demo 运行路径未测试（Bug 2/3/5 未被发现）',
    'numpy 2.x 兼容性未测试（Bug 3 未被发现）',
    'SemiAsync 跨轮陈旧度场景未测试（Bug 6 未被发现）',
    '_algo_engine 缓存的跨轮复用未测试（Bug 10 未被发现）',
]
for gap in gaps:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(gap)
    r.font.size = Pt(10.5)
    set_cjk_font(r)

# 七、修复建议
add_heading('七、修复优先级建议', 1)
add_para('按阻断 -> 严重 -> 一般的顺序修复，具体建议如下：', indent=True)
priorities = [
    'P0（振鑫）：Bug 1 - 将 ClientUpdate 增加 weights 字段（List[np.ndarray]），TrainingScheduler.aggregate() 透传真实梯度，删除 _dummy_weight_template()',
    'P0（振鑫）：Bug 2 - 补充 src/models.py、src/utils.py、src/scheduler.py（重导出）',
    'P0（振鑫）：Bug 3 - np.zeros_like(w) 改为 np.zeros(w.shape, dtype=w.dtype)',
    'P1（振鑫）：Bug 4 - ClientUpdate.weights_summary 类型改为 List[np.ndarray] 或新增 weights 字段',
    'P1（振鑫）：Bug 5 - 将 import numpy as np 移到文件顶部',
    'P1（振鑫）：Bug 6 - simulate_concurrent_round 中为不同客户端传入不同的 client_round 模拟真实陈旧度',
    'P2（建鸿）：Bug 7 - start_round 用 Pydantic BaseModel 接收请求体',
    'P2（建鸿）：Bug 8 - requirements.txt 补充 numpy',
    'P2（振鑫）：Bug 9 - SemiAsync 聚合时增加样本数加权',
    'P2（振鑫）：Bug 10 - _get_algo_engine 每轮重建或按 n_clients 做缓存键',
]
for p_text in priorities:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(p_text)
    r.font.size = Pt(10.5)
    set_cjk_font(r)

output_path = 'training-coordinator-代码审查报告.docx'
doc.save(output_path)
print('已生成：' + output_path)
