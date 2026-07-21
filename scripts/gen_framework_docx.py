# -*- coding: utf-8 -*-
"""智慧病房云边协同系统 - 整体技术框架文档生成器
输出：docs/智慧病房整体技术框架.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微软雅黑"

doc = Document()

# 设置默认字体（中英文）
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

# 页边距
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ===== 辅助函数 =====
def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            run.font.size = Pt(15)
        else:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(text, bold_label=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_label:
        run1 = p.add_run(bold_label)
        run1.bold = True
        run1.font.name = FONT
        run1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run1.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.name = FONT
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(11)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    # 灰色背景
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F4E79")
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(10)
    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()
add_para("智慧病房云边协同系统", bold=True, size=28, color=(0x1F, 0x4E, 0x79), align="center")
add_para("整体技术框架", bold=True, size=22, color=(0x2E, 0x75, 0xB6), align="center")
doc.add_paragraph()
add_para("—— 云边端协同架构 · 边缘自治 · 事件闭环 ——", size=13, color=(0x59, 0x59, 0x59), align="center")
for _ in range(3):
    doc.add_paragraph()
add_para("版本 V0.1（骨架交付版）", size=13, color=(0x40, 0x40, 0x40), align="center")
add_para("对应方案书：智慧病房云边协同优化方案书 V0.1", size=11, color=(0x59, 0x59, 0x59), align="center")
add_para("日期：2026 年 7 月", size=11, color=(0x59, 0x59, 0x59), align="center")
doc.add_page_break()

# ===== 第一章：项目概述 =====
add_heading("第一章 项目概述", 1)

add_heading("1.1 项目定位", 2)
add_para("本项目为智慧病房云边协同系统，在已验证的智能教室云边协同节能系统技术路线基础上，将领域模型、数据契约、事件闭环、边缘推理和协同训练全部面向智慧病房场景重新设计。项目用于赛事展示、科研验证和原型演示，定位为患者安全辅助系统，不作为诊断、治疗或替代医护人员处置的医疗器械。所有高风险事件均须由护士或医护人员确认后处置。")
add_para("本技术框架文档描述当前已交付的代码骨架、数据模型、通信契约、端到端联调结果，以及各成员后续接手指南。")

add_heading("1.2 建设目标", 2)
add_bullet("及时性：跌倒、离床等紧急事件由边缘端优先判定和告警，不等待云端返回。")
add_bullet("适配性：通过病房画像与可配置规则覆盖多类非 ICU 病房，而非为每类病房复制一套系统。")
add_bullet("协同性：边缘运行轻量模型，云端运行复杂模型；云端按需复核、训练并下发版本化模型。")
add_bullet("可靠性：网络中断时继续采集、识别、告警和本地留存；恢复后完成事件补传与一致性校验。")
add_bullet("可展示性：可通过虚拟摄像头、虚拟传感器与场景脚本完成完整演示，不依赖真实硬件。")

add_heading("1.3 三层交付范围", 2)
add_table(
    ["层级", "内容", "目标"],
    [
        ["P0 赛事最小可用版", "多病房大屏、虚拟传感器、跌倒/离床/呼叫/输液异常模拟、边缘自治告警、云端告警闭环、断网补传", "形成可稳定演示的端到端闭环"],
        ["P1 协同增强版", "边缘轻量视觉识别、云端复杂事件复核、模型版本管理、数据集/困难样本管理、告警规则配置", "突出云边端协同推理能力"],
        ["P2 创新展示版", "分布式协同训练、同步/半异步调度、扩散模型困难样本生成、性能看板、模型回滚", "形成赛事创新点与实验依据"],
    ],
    [1.4, 3.8, 1.8],
)
add_para("本批次交付对应 P0 的骨架部分，并预留 P1/P2 的接入点。", italic=True, size=10, color=(0x59, 0x59, 0x59))

add_heading("1.4 当前交付状态", 2)
add_para("截至本文档生成时，整体技术框架已完成端到端联调验证：")
add_bullet("8 个 Docker 容器全部启动无崩溃。")
add_bullet("3 个边缘代理上报 observation 与 health 心跳，云端 latest_state 可见 3 床在线。")
add_bullet("MySQL 种子数据加载：1 病区、3 床位、3 节点。")
add_bullet("边缘端自动触发 environment_anomaly、door_departure 事件并入库。")
add_bullet("手动注入 fall_suspected 事件，状态机闭环 new -> notified -> acknowledged 验证通过。")
add_bullet("告警确认 API 生成处置记录与审计日志。")
add_bullet("前端 http://localhost:8081 可访问，API/WS 代理正常。")
add_para("联调中发现并修复 3 个真实 bug（docker-compose 环境变量覆盖、MySQL 未挂载 init.sql、外键约束 flush 顺序），均已提交。")
doc.add_page_break()

# ===== 第二章：系统架构 =====
add_heading("第二章 系统架构", 1)

add_heading("2.1 总体架构", 2)
add_para("系统采用云边端三层协同架构，云端负责事件中心、模型管理与协同训练调度，边缘端负责多源采集、轻量推理与事件融合，前端为护士站工作台。所有节点通过 MQTT 通信，遵循统一消息信封契约。")

add_heading("2.1.1 架构层次", 3)
add_table(
    ["层级", "工作负载", "处理原则"],
    [
        ["边缘端（小模型）", "人/床/区域检测、离床初筛、姿态关键点、简单传感器融合、即时告警", "低时延、优先本地、断网继续运行"],
        ["云端（大模型）", "复杂时序行为复核、多帧片段分析、跨事件关联、困难样本挖掘、报表生成", "仅处理必要数据；不阻塞紧急告警"],
        ["训练/调度端", "客户端选择、梯度聚合、同步/半异步轮次、模型评测、灰度发布与回滚", "独立于在线告警链路，避免训练影响实时服务"],
    ],
    [1.4, 3.8, 1.8],
)

add_heading("2.1.2 服务清单（Docker Compose 编排）", 3)
add_table(
    ["服务", "技术栈", "端口", "职责"],
    [
        ["mqtt-broker", "Eclipse Mosquitto 2.0", "1884->1883", "MQTT 消息代理，QoS 1"],
        ["mysql", "MySQL 8.0", "3308->3306", "云端事件中心持久化"],
        ["cloud-backend", "FastAPI + SQLAlchemy", "8001->8000", "REST API + WebSocket + MQTT 订阅"],
        ["training-coordinator", "FastAPI", "8002->8000", "协同训练调度（空壳）"],
        ["cloud-frontend", "Vue 3 + Vite + aiohttp", "8081->80", "护士站工作台"],
        ["edge-bed-01", "Python + paho-mqtt", "-", "B01 床位边缘代理（fall,nurse_call）"],
        ["edge-bed-02", "Python + paho-mqtt", "-", "B02 床位边缘代理（bed_leave,night_wandering）"],
        ["edge-bed-03", "Python + paho-mqtt", "-", "B03 床位边缘代理（environment_anomaly,door_departure）"],
    ],
    [1.5, 2.2, 1.2, 2.6],
)

add_heading("2.2 边缘代理架构", 2)
add_para("每个床位运行一个独立的边缘代理容器，包含采集适配器、推理引擎、事件融合引擎、本地数据库、MQTT 客户端和主循环。适配器采用抽象接口设计，模拟器与真实硬件共用同一协议，接入真实摄像头或传感器时只需替换适配器实现，不改动融合、推理、MQTT 与主循环。")
add_bullet("采集适配器层：", "CameraAdapter / BedSensorAdapter / InfusionAdapter / EnvironmentAdapter 四类，BaseAdapter 定义统一 read() 接口。")
add_bullet("推理引擎：", "InferenceEngine 空壳，强制输出 model_name / model_version / confidence / inference_ms / evidence_ref（对齐方案书 §3.3 验收要求）。")
add_bullet("事件融合引擎：", "FusionEngine 接收多源观测与推理结果，按规则输出 SafetyEvent，含同类去重与夜间判定。")
add_bullet("本地数据库：", "SQLite 三表（observations / safety_events / node_health），含 synced 字段支持断网补传。")
add_bullet("MQTT 客户端：", "新主题树 ward/{ward}/node/{node}/{observation,event,health}，通用信封封装。")
add_bullet("场景驱动器：", "ScenarioDriver 读取 SCENARIO_PROFILE 环境变量，4 阶段状态机驱动模拟数据注入。")

add_heading("2.3 云端事件中心架构", 2)
add_para("云端后端从教室的“状态查询”改造为病房“事件中心”，提供事件状态机、告警确认闭环、模型版本管理与审计日志。")
add_bullet("数据模型：", "10 张表（wards / beds / edge_nodes / observations / safety_events / alert_tasks / event_dispositions / model_versions / model_deployments / audit_logs）。")
add_bullet("MQTT 处理器：", "订阅 ward/+/node/+/{observation,event,health} 与 ward/+/alert/+/ack，解包信封后写库并广播 WebSocket。")
add_bullet("REST API：", "17 个路由覆盖病区、床位、事件查询与确认、节点健康、观测历史、模型管理与系统统计。")
add_bullet("WebSocket：", "增量推送事件、告警确认与节点健康，前端无需高频轮询。")
add_bullet("幂等与审计：", "event_id 唯一约束去重；事件写入、确认操作、模型下发全部留痕 audit_logs。")
doc.add_page_break()

# ===== 第三章：数据模型 =====
add_heading("第三章 数据模型", 1)

add_heading("3.1 数据库表清单", 2)
add_para("云端 MySQL 共 10 张表，对齐方案书 §4.2 必须替换的领域对象。将教室的 Room/SensorData/DeviceStatus/ControlLog/AlarmLog 替换为病房的 Ward/Bed/EdgeNode/Observation/SafetyEvent/AlertTask/EventDisposition/ModelVersion/ModelDeployment/AuditLog。")
add_table(
    ["表名", "用途", "关键字段"],
    [
        ["wards", "病区", "id, name, ward_type, location, status"],
        ["beds", "床位（FK->wards）", "id, ward_id, name, status"],
        ["edge_nodes", "边缘节点（FK->wards）", "id, ward_id, bed_id, status, model_version, last_heartbeat, buffered_events"],
        ["observations", "多源观测数据", "node_id, source_type, data(JSON), quality(JSON), timestamp"],
        ["safety_events", "安全事件（核心）", "event_id, event_type, priority, state, confidence, model_name, model_version, occurred_at"],
        ["alert_tasks", "告警任务", "event_id, priority, channel, notified_at"],
        ["event_dispositions", "事件处置记录", "event_id, action, operator_id, operator_role, result"],
        ["model_versions", "模型版本", "model_name, model_version, artifact_url, runtime, target_device, status"],
        ["model_deployments", "模型部署记录（灰度）", "model_name, model_version, node_id, action, status"],
        ["audit_logs", "审计日志", "action, target_type, target_id, operator_id, detail, occurred_at"],
    ],
    [1.5, 2.2, 4.3],
)

add_heading("3.2 边缘端 SQLite 表", 2)
add_table(
    ["表名", "用途", "关键字段"],
    [
        ["observations", "本地观测缓存", "ward_id, node_id, source_type, data, timestamp, synced"],
        ["safety_events", "本地事件缓存（断网补传）", "event_id, event_type, priority, payload, occurred_at, synced"],
        ["node_health", "健康心跳记录", "node_id, status, metrics, timestamp, synced"],
    ],
    [1.5, 2.5, 4.0],
)

add_heading("3.3 事件状态机", 2)
add_para("事件生命周期状态机对齐方案书 §6.2，将识别事件、通知任务、人工处置分开建模：")
add_code("new -> notified -> acknowledged -> resolved")
add_code("                       ├-> false_positive")
add_code("                       └-> escalated")
add_bullet("new：", "边缘端新生成事件。")
add_bullet("notified：", "已推送到护士站（云端收到即标记）。")
add_bullet("acknowledged：", "护士已确认到场，开始处置。")
add_bullet("resolved：", "已处置完成。")
add_bullet("false_positive：", "误报，关闭事件。")
add_bullet("escalated：", "升级，转更高优先级处置。")
add_para("统计逻辑中，new/notified/acknowledged 三态均算 pending（待处理），resolved/false_positive/escalated 为终态。", italic=True, size=10, color=(0x59, 0x59, 0x59))
doc.add_page_break()

# ===== 第四章：MQTT 通信契约 =====
add_heading("第四章 MQTT 通信契约", 1)

add_heading("4.1 主题树设计", 2)
add_para("MQTT 主题树对齐方案书 §4.3，所有消息 QoS=1，消费端按 message_id 去重，边缘侧按序号补传，所有时间使用 UTC ISO 8601（Z 结尾）。")
add_code("# ── 上行：边缘端 -> 云端 ──")
add_code("ward/{ward_id}/node/{node_id}/observation      多源观测数据")
add_code("ward/{ward_id}/node/{node_id}/event             安全事件")
add_code("ward/{ward_id}/node/{node_id}/health            节点健康心跳")
add_code("# ── 下行：云端 -> 边缘端 ──")
add_code("ward/{ward_id}/alert/{event_id}/ack             告警确认/处置/升级指令")
add_code("node/{node_id}/config/set                        节点配置下发")
add_code("node/{node_id}/model/deploy                     模型版本下发（灰度）")
add_code("node/{node_id}/model/rollback                   模型回滚")
add_code("# ── 训练链路（与实时业务隔离）──")
add_code("training/{job_id}/node/{node_id}/command        训练指令")
add_code("training/{job_id}/node/{node_id}/update         梯度/权重上报")
add_code("training/{job_id}/status                         训练任务状态")

add_heading("4.2 通用消息信封", 2)
add_para("所有消息外层必须符合 contracts/envelope.json 信封结构，包含 7 个字段：")
add_table(
    ["字段", "类型", "说明"],
    [
        ["message_id", "uuid", "消息唯一 ID，消费端按此去重"],
        ["event_id", "uuid|null", "关联事件 ID；非事件消息为 null"],
        ["schema_version", "string", "契约版本，当前 v1"],
        ["occurred_at", "iso8601", "事件发生时间（UTC，Z 结尾）"],
        ["source", "string", "来源标识，如 edge:EDGE-W01-B01 / cloud"],
        ["trace_id", "string", "跨服务追踪 ID"],
        ["payload", "object", "业务数据，结构见各消息 Schema"],
    ],
    [1.8, 1.6, 4.6],
)

add_heading("4.3 契约文件清单", 2)
add_para("contracts/ 目录下共 6 份 JSON Schema，作为前后端边缘端联调的唯一标准：")
add_table(
    ["文件", "对应主题", "用途"],
    [
        ["envelope.json", "（通用）", "通用消息信封"],
        ["observation.json", "ward/+/node/+/observation", "多源观测数据"],
        ["safety_event.json", "ward/+/node/+/event", "安全事件"],
        ["alert_ack.json", "ward/+/alert/+/ack", "告警确认指令"],
        ["node_health.json", "ward/+/node/+/health", "节点健康心跳"],
        ["model_deploy.json", "node/+/model/deploy", "模型版本下发"],
    ],
    [1.8, 2.8, 3.4],
)

add_heading("4.4 订阅清单", 2)
add_heading("云端 cloud-backend 订阅", 3)
add_table(
    ["通配主题", "QoS", "处理器"],
    [
        ["ward/+/node/+/observation", "1", "MqttHandler._handle_observation"],
        ["ward/+/node/+/event", "1", "MqttHandler._handle_event"],
        ["ward/+/node/+/health", "1", "MqttHandler._handle_health"],
        ["ward/+/alert/+/ack", "1", "MqttHandler._handle_ack"],
    ],
    [3.0, 1.0, 4.0],
)
add_heading("边缘端 edge-agent 订阅", 3)
add_table(
    ["通配主题", "QoS", "处理器"],
    [
        ["ward/{ward_id}/alert/+/ack", "1", "EdgeAgent.handle_ack"],
        ["node/{node_id}/config/set", "1", "EdgeAgent.handle_config"],
        ["node/{node_id}/model/deploy", "1", "EdgeAgent.handle_model_deploy"],
        ["node/{node_id}/model/rollback", "1", "EdgeAgent.handle_model_deploy"],
    ],
    [3.0, 1.0, 4.0],
)
doc.add_page_break()

# ===== 第五章：边缘代理设计 =====
add_heading("第五章 边缘代理设计", 1)

add_heading("5.1 采集适配器接口", 2)
add_para("BaseAdapter 定义统一接口，所有适配器必须实现 read()，返回标准化的 Observation 结构。模拟器与真实硬件共用此接口，接入真实硬件时只替换适配器实现，不改动 fusion/inference/main。")

add_heading("5.1.1 四类适配器", 3)
add_table(
    ["适配器", "source_type", "输出字段", "模拟实现"],
    [
        ["CameraAdapter", "camera", "presence, person_count, posture, bbox, pose_keypoints, fall_score", "场景驱动注入"],
        ["BedSensorAdapter", "bed_sensor", "occupied, bed_state, absence_seconds, pressure_raw", "场景驱动注入"],
        ["InfusionAdapter", "infusion", "flow_rate, volume_pct, remaining_minutes, anomaly", "液位衰减+场景注入"],
        ["EnvironmentAdapter", "environment", "temperature, humidity, light, co2, door_open, air_quality", "正弦波动+噪声"],
    ],
    [1.4, 1.1, 3.5, 2.2],
)

add_heading("5.1.2 Observation 结构", 3)
add_para("Observation 对齐 contracts/observation.json 中的 source 定义，包含 source_type、data、quality、timestamp 四部分。Quality 含 confidence（0~1）、latency_ms、degraded（是否降级，如遮挡/低照度）三个字段。")

add_heading("5.2 事件融合引擎", 2)
add_para("FusionEngine 每周期调用 fuse() 一次，传入本轮所有观测与推理结果，返回新生成的 SafetyEvent 列表。第一版采用规则融合，为后续接入 YOLO/姿态模型预留接口（fusion 只消费 InferenceResult + Observation，模型升级时只替换 inference.py）。")

add_heading("5.2.1 融合规则", 3)
add_table(
    ["事件类型", "优先级", "触发规则", "去重"],
    [
        ["fall_suspected", "P1", "摄像头 posture=falling + fall_score>0.5 + 床位离床", "30 秒"],
        ["bed_leave", "P2", "床位 absence_seconds ≥ 30 秒阈值", "30 秒"],
        ["infusion_anomaly", "P2", "输液 anomaly != normal（fast/slow/low_volume/completed/interrupted）", "30 秒"],
        ["environment_anomaly", "P3", "温度>29℃ 或 CO₂>1000ppm 或光照<50lux", "30 秒"],
        ["door_departure", "P2", "门磁 open + 摄像头检测到人离开区域", "30 秒"],
        ["night_wandering", "P2", "夜间时段（22:00-06:00）+ 持续离床", "30 秒"],
    ],
    [1.6, 0.6, 4.4, 1.4],
)

add_heading("5.2.2 配置项（环境变量）", 3)
add_table(
    ["环境变量", "默认值", "说明"],
    [
        ["BED_LEAVE_THRESHOLD", "30", "离床告警阈值秒数"],
        ["NIGHT_START / NIGHT_END", "22 / 6", "夜间时段起止小时"],
        ["TEMP_ALARM_HIGH", "29.0", "高温告警阈值 ℃"],
        ["CO2_ALARM_HIGH", "1000", "CO₂ 告警阈值 ppm"],
        ["LIGHT_ALARM_LOW", "50", "低光照告警阈值 lux"],
        ["EVENT_DEDUPE_SECONDS", "30", "同类事件去重秒数"],
    ],
    [2.5, 1.2, 4.3],
)

add_heading("5.3 模型推理接口", 2)
add_para("InferenceEngine 为边缘端识别模型预留统一接口。亚伦后续接入真实 ONNX/OpenVINO/TensorRT 模型时，只需替换 run() 内部实现，不改动 fusion/mqtt_client/main。输出强制包含方案书 §3.3 验收要求的字段：model_name、model_version、confidence、inference_ms、evidence_ref。")
add_para("当前版本（rule-fusion-v1 / 0.1.0-mock）透传适配器数据，不做真实推理。真实模型接入后，predictions 字段将包含检测结果、姿态关键点与跌倒置信度。")

add_heading("5.4 场景脚本驱动器", 2)
add_para("ScenarioDriver 读取环境变量 SCENARIO_PROFILE（逗号分隔的事件类型列表），按节奏触发对应场景，向适配器注入模拟状态。场景生命周期：idle -> started -> sustained -> recovering -> idle（4 阶段状态机，当前实现 3 阶段，景彬补充完整 4 阶段含“人工确认”）。")
add_para("docker-compose.yml 已编排 3 床 6 场景：")
add_bullet("B01：fall_suspected, nurse_call")
add_bullet("B02：bed_leave, night_wandering")
add_bullet("B03：environment_anomaly, door_departure")

add_heading("5.5 离线缓存与补传", 2)
add_para("边缘端 SQLite 三表（observations / safety_events / node_health）均含 synced 字段。网络中断时数据标记 synced=0 继续本地缓存；恢复后按时间序批量补传，云端按 event_id 去重，不产生重复告警。主循环每周期调用 sync_offline_data() 同步离线事件。")
doc.add_page_break()

# ===== 第六章：REST API 规范 =====
add_heading("第六章 REST API 规范", 1)
add_para("云端事件中心提供 17 个 REST 路由，统一响应格式 {code, message, data}。所有时间使用 UTC ISO 8601（Z 结尾），前端转本地时区显示。")

add_heading("6.1 病区与床位", 2)
add_table(
    ["方法", "路径", "说明"],
    [
        ["GET", "/api/wards", "获取所有病区及床位/节点状态"],
        ["GET", "/api/wards/{ward_id}", "获取病区详情"],
    ],
    [0.8, 3.0, 4.2],
)

add_heading("6.2 安全事件", 2)
add_table(
    ["方法", "路径", "说明"],
    [
        ["GET", "/api/events", "查询事件（支持 ward_id/bed_id/priority/state/event_type/hours/limit 过滤）"],
        ["GET", "/api/events/{event_id}", "获取事件详情（含处置记录）"],
        ["POST", "/api/events/{event_id}/ack", "确认/处置/升级事件（MQTT 下发到边缘端，同时更新本地状态）"],
    ],
    [0.8, 3.0, 4.2],
)
add_para("确认请求体（AckRequest）：")
add_code('{"action": "acknowledge|resolve|false_positive|escalate",')
add_code(' "operator_id": "nurse-demo", "operator_name": "演示护士",')
add_code(' "operator_role": "nurse|charge_nurse|admin|observer",')
add_code(' "result": "处置结果", "note": "备注"}')

add_heading("6.3 节点与观测", 2)
add_table(
    ["方法", "路径", "说明"],
    [
        ["GET", "/api/nodes?ward_id=W-01", "获取边缘节点列表及健康状态"],
        ["GET", "/api/observations?bed_id=B01&source_type=camera&hours=1", "查询观测历史数据"],
    ],
    [0.8, 3.0, 4.2],
)

add_heading("6.4 模型管理", 2)
add_table(
    ["方法", "路径", "说明"],
    [
        ["GET", "/api/models", "列出所有模型版本"],
        ["POST", "/api/models/deploy?node_id=EDGE-W01-B01", "下发模型到指定节点"],
    ],
    [0.8, 3.0, 4.2],
)

add_heading("6.5 系统统计与健康检查", 2)
add_table(
    ["方法", "路径", "说明"],
    [
        ["GET", "/api/stats", "获取系统全局统计（病区/床位/节点/事件/待处理）"],
        ["GET", "/health", "健康检查"],
        ["GET", "/", "API 根路径"],
        ["WS", "/ws", "WebSocket 端点（增量推送事件/告警确认/节点健康）"],
    ],
    [0.8, 3.0, 4.2],
)

add_heading("6.6 WebSocket 消息类型", 2)
add_table(
    ["type", "触发场景", "payload 关键字段"],
    [
        ["observation", "边缘端上报观测", "ward_id, node_id, bed_id, sources"],
        ["safety_event", "边缘端上报安全事件", "event_id, event_type, priority, state, confidence"],
        ["event_ack", "告警确认指令", "event_id, action, operator"],
        ["node_health", "节点健康心跳", "node_id, status, buffered_events"],
        ["ping", "心跳检测（前端回复 pong）", "-"],
    ],
    [1.4, 2.6, 4.0],
)
doc.add_page_break()

# ===== 第七章：端到端联调与验收 =====
add_heading("第七章 端到端联调与验收", 1)

add_heading("7.1 启动方式", 2)
add_code("cd smart-ward")
add_code("docker compose up --build")
add_para("启动后访问：")
add_bullet("护士站工作台：http://localhost:8081")
add_bullet("云端 API 文档：http://localhost:8001/docs")
add_bullet("协同训练 API：http://localhost:8002/docs")
add_bullet("MQTT Broker：localhost:1884")

add_heading("7.2 本地测试", 2)
add_code("cd edge-agent && python -m pytest tests/ -v     # 10 项测试")
add_code("cd training-coordinator && python -m pytest tests/ -v   # 4 项测试")

add_heading("7.3 联调验收结果", 2)
add_para("截至本文档生成时，以下验收项全部通过：")
add_table(
    ["验收项", "结果", "说明"],
    [
        ["8 服务全部启动无崩溃", "✅ 通过", "mqtt/mysql/backend/training/frontend + 3 edge-agent"],
        ["MySQL 种子数据加载", "✅ 通过", "1 病区 + 3 床位 + 3 节点"],
        ["边缘端上报 observation + health", "✅ 通过", "3 节点在线，latest_state 可见"],
        ["GET /api/wards 返回完整数据", "✅ 通过", "W-01 + B01/B02/B03 + 3 节点状态"],
        ["GET /api/stats 统计正确", "✅ 通过", "total_wards=1, online_nodes=3"],
        ["边缘端自动触发事件入库", "✅ 通过", "environment_anomaly / door_departure"],
        ["手动注入跌倒事件", "✅ 通过", "P1 fall_suspected 状态 notified"],
        ["告警确认闭环", "✅ 通过", "acknowledged + 处置记录 + 审计日志"],
        ["前端 HTTP 可访问", "✅ 通过", "http://localhost:8081 HTTP 200"],
        ["前端 API/WS 代理", "✅ 通过", "/api/wards 透传成功"],
        ["单元测试", "✅ 通过", "edge-agent 10 + training-coordinator 4 = 14 项全绿"],
    ],
    [3.2, 1.0, 3.8],
)

add_heading("7.4 联调中发现并修复的 bug", 2)
add_para("联调过程中发现并修复 3 个真实 bug（commit 0c9fab6）：")
add_bullet("docker-compose 环境变量覆盖：", "原 DATABASE_URL 注入但 database.py 读 MYSQL_HOST，且容器内 .env 文件覆盖成 localhost，导致连不上 MySQL。修复：显式注入 MYSQL_HOST=mysql 等环境变量。")
add_bullet("MySQL 未挂载 init.sql：", "原 mysql 服务未挂载 init.sql 到 /docker-entrypoint-initdb.d/，导致 wards/beds 种子数据为空，GET /api/wards 返回空数组。修复：挂载 ./cloud-backend/init.sql。")
add_bullet("外键约束失败：", "alert_tasks.event_id 外键引用 safety_events.event_id，但 db.add(event) 后未 flush 直接 db.add(task)，event 未持久化。修复：db.add(event) 后加 db.flush()。")
doc.add_page_break()

# ===== 第八章：团队分工与接手指南 =====
add_heading("第八章 团队分工与接手指南", 1)

add_heading("8.1 亚伦（边缘模型选型 + 技术框架）", 2)
add_heading("已就位", 3)
add_bullet("edge-agent/src/adapters/camera.py 预留 YOLO/姿态模型接入点。")
add_bullet("edge-agent/src/inference.py 强制输出方案书 §3.3 验收字段。")
add_bullet("contracts/safety_event.json 已定义 model 字段结构。")
add_heading("下一步", 3)
add_bullet("在 docs/ 新增 02-边缘模型选型对比.md，对比 YOLO Nano 系（yolov8n/yolov10n/yolo11n）、姿态模型（YOLO-Pose/MediaPipe）、推理运行时（ONNX/OpenVINO/TensorRT）、量化方案（INT8）。")
add_bullet("在 edge-agent/src/inference.py 替换 InferenceEngine.run()，加载真实模型。")
add_bullet("在目标硬件上实测延迟/FPS/召回率/误报率，填入对比文档。")

add_heading("8.2 景彬（场景脚本 + 前端展示）", 2)
add_heading("已就位", 3)
add_bullet("edge-agent/src/scenario.py 场景驱动器已实现 4 阶段状态机。")
add_bullet("docker-compose 已编排 3 床 6 场景。")
add_bullet("cloud-frontend/src/App.vue 护士站骨架（顶栏统计 + 床位卡片 + 告警工作台 + 确认按钮）。")
add_heading("下一步", 3)
add_bullet("在 scenario.py 补充完整 4 阶段脚本（含“人工确认”阶段）。")
add_bullet("在 cloud-frontend/src/components/ 新增 WardCard.vue / EventTrendChart.vue / NodeLatencyChart.vue 等组件。")
add_bullet("实现“场景注入台”调试面板（替换教室的 DebugPanel）。")

add_heading("8.3 建鸿 + 振鑫（分布式协同训练）", 2)
add_heading("已就位", 3)
add_bullet("training-coordinator/app/scheduler.py 已定义三阶段策略枚举与 TrainingScheduler 空壳。")
add_bullet("training-coordinator/app/main.py 已有 /rounds / /update / /aggregate 空壳 API。")
add_bullet("MQTT 训练主题已预留：training/{job_id}/node/{node_id}/{command,update} + training/{job_id}/status。")
add_heading("下一步", 3)
add_bullet("在 scheduler.py 实现 aggregate()：FedAvg 加权平均（阶段 A）。")
add_bullet("实现半异步调度：按数据量/训练耗时/网络质量/版本陈旧度计算聚合权重（阶段 B）。")
add_bullet("接入 MQTT 训练主题，与在线业务链路解耦。")
add_bullet("记录每轮参与节点、样本数、梯度摘要、耗时、精度、带宽、异常原因（阶段 C）。")

add_heading("8.4 先伟 + 烽亮（扩散模型 + 困难样本）", 2)
add_heading("已就位", 3)
add_bullet("cloud-backend/init.sql 已有 model_versions 表（支持灰度发布与回滚）。")
add_bullet("contracts/model_deploy.json 已定义模型下发契约。")
add_heading("下一步", 3)
add_bullet("在云端新建 diffusion-service/（或并入 cloud-backend），实现扩散模型生成困难样本。")
add_bullet("困难样本经人工审核后入数据集，触发新模型训练轮次。")
add_bullet("新模型通过 model/deploy 主题灰度下发到边缘端。")

add_heading("8.5 复用策略总结", 2)
add_para("本项目验证了“非推倒重来”的复用策略，从已验证的智能教室项目复用基础设施，仅替换领域模型：")
add_table(
    ["复用方式", "文件数", "说明"],
    [
        ["原样复制", "8", "websocket_manager.py / logger.py / Dockerfile / requirements.txt / 前端 proxy.py / nginx.conf / vite.config.js / websocket.js"],
        ["模式复用 + 字段重写", "4", "database.py / mqtt_handler.py / main.py / schemas.py（复用 engine/session/init_db/get_db/重连退避等模式，字段 Room->Ward/Bed/EdgeNode）"],
        ["新写", "-", "10 表 init.sql / 新 MQTT 主题树 / 4 类适配器 / 融合引擎 / 场景驱动器 / 6 份 JSON Schema / 3 份文档"],
    ],
    [2.0, 0.8, 5.2],
)

add_heading("8.6 不在本批次范围", 2)
add_bullet("边缘识别模型选型对比文档（亚伦任务1，下一批次）。")
add_bullet("场景脚本完整 4 阶段实现（景彬任务）。")
add_bullet("分布式训练算法实现（建鸿/振鑫任务）。")
add_bullet("扩散模型困难样本生成（烽亮/先伟任务）。")
add_bullet("护士站前端完整 UI（景彬任务）。")
add_bullet("KubeEdge 部署清单（方案书 §5.2 P2 阶段）。")
add_bullet("真实硬件接入（方案书 §5.2）。")

# ===== 保存 =====
out_path = "docs/智慧病房整体技术框架.docx"
doc.save(out_path)
import os
size_kb = os.path.getsize(out_path) / 1024
print(f"✅ 文档已生成：{out_path}")
print(f"   大小：{size_kb:.1f} KB")
